import streamlit as st
import pandas as pd
import numpy as np
import pulp
import pydeck as pdk
import json
import warnings
import matplotlib.pyplot as plt
import seaborn as sns
import os
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

import math

warnings.filterwarnings('ignore')
sns.set_theme(style="whitegrid")

# --- НАСТРОЙКА СТРАНИЦЫ ---
st.set_page_config(page_title="ГИС Оптимизация ТСО v5.5", layout="wide")
st.markdown(
    """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stAppDeployButton {display:none;}
    header {visibility: hidden;}
    </style>
    """,
    unsafe_allow_html=True
)
# --- ИНИЦИАЛИЗАЦИЯ БАЗОВОГО КАТАЛОГА В ПАМЯТИ ---
DEFAULT_CATALOG = [
    {"name": "Речевые уст.", "ch": "Проводной", "cost": 100, "cov": 2500, "rel": 0.95, "time": 10, "k_act": 0.90},
    {"name": "Речевые уст.", "ch": "Сотовая", "cost": 120, "cov": 2500, "rel": 0.95, "time": 10, "k_act": 0.90},
    {"name": "Речевые уст.", "ch": "IP", "cost": 80, "cov": 2500, "rel": 0.95, "time": 5, "k_act": 0.90},
    {"name": "Мобильные комп.", "ch": "Сотовая", "cost": 200, "cov": 3000, "rel": 0.98, "time": 20, "k_act": 0.85},
    {"name": "Мобильные комп.", "ch": "Спутник", "cost": 1500, "cov": 3000, "rel": 0.98, "time": 25, "k_act": 0.85},
    {"name": "SMS-оповещение", "ch": "Сотовая", "cost": 50, "cov": 10000, "rel": 0.95, "time": 5, "k_act": 0.80},
    {"name": "Моб. приложения", "ch": "Сотовая", "cost": 100, "cov": 10000, "rel": 0.95, "time": 2, "k_act": 0.80},
    {"name": "Моб. приложения", "ch": "IP", "cost": 30, "cov": 10000, "rel": 0.95, "time": 2, "k_act": 0.80},
    {"name": "ТВ-системы", "ch": "Спутник", "cost": 2000, "cov": 10000, "rel": 0.95, "time": 15, "k_act": 0.40},
    {"name": "ТВ-системы", "ch": "ТВ/радио", "cost": 1000, "cov": 10000, "rel": 0.95, "time": 15, "k_act": 0.40},
    {"name": "Радиовещание", "ch": "Радио", "cost": 150, "cov": 8000, "rel": 0.95, "time": 10, "k_act": 0.50},
    {"name": "Радиовещание", "ch": "Спутник", "cost": 1800, "cov": 8000, "rel": 0.95, "time": 15, "k_act": 0.50},
    {"name": "Радиовещание", "ch": "ТВ/радио", "cost": 800, "cov": 8000, "rel": 0.95, "time": 15, "k_act": 0.50},
    {"name": "БАС (Дроны)", "ch": "Сотовая", "cost": 300, "cov": 4000, "rel": 0.95, "time": 15, "k_act": 0.85},
    {"name": "БАС (Дроны)", "ch": "Спутник", "cost": 450, "cov": 4000, "rel": 0.95, "time": 20, "k_act": 0.85}
]

if 'catalog' not in st.session_state:
    st.session_state.catalog = [dict(item) for item in DEFAULT_CATALOG]


# ==========================================
# ФУНКЦИЯ ИИ-АНАЛИТИКИ
# ==========================================
def generate_ai_insights(total_cost, total_coverage, total_population, summary_df):
    api_key = st.secrets.get("OPENROUTER_API_KEY", os.environ.get("OPENROUTER_API_KEY"))
    if not api_key:
        return "Ошибка: API-ключ OpenRouter не найден. Убедитесь, что вы добавили его в конфигурацию среды."

    try:
        client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)
        coverage_percent = (total_coverage / total_population * 100) if total_population > 0 else 0
        data_context = summary_df[['Система_и_Канал', 'Количество', 'Общая_стоимость', 'Общий_охват']].to_string(index=False)

        prompt = f"""
        Вы — Главный системный аналитик и руководитель проектного офиса по модернизации систем гражданской обороны (РАСЦО).

        ВВОДНЫЕ ДАННЫЕ ПРОЕКТА (ОБЯЗАТЕЛЬНО ИСПОЛЬЗУЙТЕ ЭТИ ЦИФРЫ В ТЕКСТЕ):
        - Итоговый бюджет распределения: {total_cost:,.0f} у.е.
        - Расчетный охват населения: {total_coverage:,.0f} чел. ({coverage_percent:.1f}% от всей зоны риска).
        - Детализация по оборудованию:\n{data_context}

        ЗАДАЧА:
        Сформируйте официальную аналитическую записку.
        КРИТИЧЕСКОЕ ТРЕБОВАНИЕ 1: КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО использовать Markdown-разметку (никаких звездочек **, решеток ###, жирного шрифта). Используйте строго обычный текст.
        КРИТИЧЕСКОЕ ТРЕБОВАНИЕ 2: Ваш отчет должен быть насыщен реальными цифрами из Вводных данных (упоминайте точные суммы бюджета, количество устройств, охват в людях и проценты), чтобы анализ выглядел глубоким и математически обоснованным.

        СТРУКТУРА:
        1. ВВЕДЕНИЕ: Обоснование применения симплекс-метода и расчета мультириска.
        [ТАБЛИЦА_ОБОРУДОВАНИЯ] (напишите этот тег ровно в этом месте)
        2. АНАЛИЗ ЭФФЕКТИВНОСТИ: Научное объяснение выбора каналов связи с точки зрения рентабельности. (Обязательно приведите цифры стоимости и охвата конкретных ТСО из таблицы).
        [ГРАФИК_РАСПРЕДЕЛЕНИЯ] (напишите этот тег ровно в этом месте)
        [ГРАФИК_ОХВАТА] (напишите этот тег ровно в этом месте)
        3. РЕКОМЕНДАЦИИ ПО ИНТЕГРАЦИИ: 3-4 конкретных шага по внедрению и эксплуатации системы.
        """

        response = client.chat.completions.create(
            model="openai/gpt-4o",
            messages=[
                {"role": "system", "content": "Вы строгий и профессиональный государственный аналитик. Не используете эмоции и спецсимволы в тексте. Оперируете только сухими цифрами и фактами."},
                {"role": "user", "content": prompt}
            ],
            extra_headers={
                "HTTP-Referer": "https://github.com/",
                "X-Title": "GIS Warning Optimizer"
            }
        )
        # Принудительная очистка текста от любых оставшихся Markdown символов
        clean_text = response.choices[0].message.content.replace('*', '').replace('#', '')
        return clean_text
    except Exception as e:
        return f"Ошибка соединения с ИИ: {e}"


# --- 1. ЗАГРУЗКА ГРАНИЦ ТАТАРСТАНА ---
@st.cache_data
def get_tatarstan_geojson():
    try:
        with open('tatarstan_districts_osm.geojson', 'r', encoding='utf-8') as f:
            data = json.load(f)
        for feature in data.get('features', []):
            geom = feature.get('geometry', {})
            if geom.get('type') == 'Polygon':
                new_coords = [[[pt[1], pt[0]] if pt[0] > pt[1] else pt for pt in ring] for ring in
                              geom.get('coordinates', [])]
                geom['coordinates'] = new_coords
            elif geom.get('type') == 'MultiPolygon':
                new_coords = [[[[pt[1], pt[0]] if pt[0] > pt[1] else pt for pt in ring] for ring in poly] for poly in
                              geom.get('coordinates', [])]
                geom['coordinates'] = new_coords
        return data
    except Exception as e:
        return None


# --- 2. КЭШИРОВАНИЕ И ИНТЕГРАЦИЯ ДАННЫХ ---


def _first_non_empty_value(series, default=''):
    """Возвращает наиболее частое непустое значение из Series."""
    try:
        cleaned = series.dropna().astype(str)
        cleaned = cleaned[cleaned.str.strip() != '']
        if cleaned.empty:
            return default
        return cleaned.mode().iloc[0]
    except Exception:
        return default


def _safe_max_date(series, default='Нет данных'):
    """Безопасно возвращает максимальную дату из Series со смешанными типами."""
    try:
        dt = pd.to_datetime(series, errors='coerce')
        dt = dt.dropna()
        if dt.empty:
            return default
        return dt.max()
    except Exception:
        return default


def aggregate_nearby_threat_zones(df_zones, radius_m=500):
    """
    Объединяет опасные зоны, центры которых находятся в радиусе radius_m друг от друга.
    Для каждой группы оставляется одна усредненная зона, чтобы далее подобрать одну ТСО
    той же MILP-моделью run_optimization.
    """
    if df_zones is None or df_zones.empty:
        return df_zones

    df = df_zones.copy().reset_index(drop=True)
    df['lat_cluster'] = pd.to_numeric(df['lat_cluster'], errors='coerce')
    df['lon_cluster'] = pd.to_numeric(df['lon_cluster'], errors='coerce')
    df = df.dropna(subset=['lat_cluster', 'lon_cluster']).reset_index(drop=True)

    if len(df) <= 1:
        df['Кол_во_опасностей_в_группе'] = len(df)
        return df

    lat0 = 55.5
    meters_per_degree_lat = 111_320
    meters_per_degree_lon = 111_320 * math.cos(math.radians(lat0))

    xs = df['lon_cluster'].astype(float).to_numpy() * meters_per_degree_lon
    ys = df['lat_cluster'].astype(float).to_numpy() * meters_per_degree_lat

    parent = list(range(len(df)))
    rank = [0] * len(df)

    def find(a):
        while parent[a] != a:
            parent[a] = parent[parent[a]]
            a = parent[a]
        return a

    def union(a, b):
        ra, rb = find(a), find(b)
        if ra == rb:
            return
        if rank[ra] < rank[rb]:
            parent[ra] = rb
        elif rank[ra] > rank[rb]:
            parent[rb] = ra
        else:
            parent[rb] = ra
            rank[ra] += 1

    # Быстрый пространственный индекс: сравниваем точку только с соседними ячейками.
    cell_size = float(radius_m)
    grid = {}
    for i, (x, y) in enumerate(zip(xs, ys)):
        cx = int(math.floor(x / cell_size))
        cy = int(math.floor(y / cell_size))
        for gx in range(cx - 1, cx + 2):
            for gy in range(cy - 1, cy + 2):
                for j in grid.get((gx, gy), []):
                    dx = x - xs[j]
                    dy = y - ys[j]
                    if dx * dx + dy * dy <= radius_m * radius_m:
                        union(i, j)
        grid.setdefault((cx, cy), []).append(i)

    groups = {}
    for i in range(len(df)):
        groups.setdefault(find(i), []).append(i)

    rows = []
    for group_indices in groups.values():
        g = df.iloc[group_indices].copy()

        lat_mean = float(g['lat_cluster'].mean())
        lon_mean = float(g['lon_cluster'].mean())
        row = {
            'Район': _first_non_empty_value(g['Район'], 'Нет данных') if 'Район' in g.columns else 'Нет данных',
            'Населенный_пункт': _first_non_empty_value(g['Населенный_пункт'], 'Нет данных') if 'Населенный_пункт' in g.columns else 'Нет данных',
            'lat_cluster': round(lat_mean, 6),
            'lon_cluster': round(lon_mean, 6),
            'lat_key': round(lat_mean, 2),
            'lon_key': round(lon_mean, 2),
            'acq_date': _safe_max_date(g['acq_date']) if 'acq_date' in g.columns else 'Нет данных',
            'Кол_во_инцидентов': int(pd.to_numeric(g.get('Кол_во_инцидентов', 1), errors='coerce').fillna(1).sum()),
            'Кол_во_опасностей_в_группе': int(len(g)),
            'Старая_сирена': 'НЕТ'
        }

        for col in ['Индекс_Огня', 'Индекс_Воды', 'До_ближайшей_2G_вышки_км', 'До_ближайшей_4G_вышки_км']:
            if col in g.columns:
                row[col] = float(pd.to_numeric(g[col], errors='coerce').fillna(0).mean())

        rows.append(row)

    result = pd.DataFrame(rows)

    # Гарантируем наличие служебных колонок, нужных дальнейшей модели.
    for col, default in {
        'Индекс_Огня': 0.0,
        'Индекс_Воды': 0.0,
        'До_ближайшей_2G_вышки_км': 10.0,
        'До_ближайшей_4G_вышки_км': 10.0,
        'acq_date': 'Нет данных',
        'Старая_сирена': 'НЕТ'
    }.items():
        if col not in result.columns:
            result[col] = default

    return result.reset_index(drop=True)

@st.cache_data
def load_data():
    try:
        df_matrix = pd.read_excel('СУПЕР_МАТРИЦА_ЭТАЛОН_ОБЪЕДИНЕННАЯ.xlsx')

        # Старый файл теперь используется только как источник опасностей и координат.
        # Население из него НЕ берем. Для отображения радиуса 400 м важно сохранять
        # реальные координаты опасностей, а не укрупнять их округлением до 0.01/0.02 градуса.
        df_matrix['latitude'] = pd.to_numeric(df_matrix['latitude'], errors='coerce')
        df_matrix['longitude'] = pd.to_numeric(df_matrix['longitude'], errors='coerce')
        df_matrix = df_matrix.dropna(subset=['latitude', 'longitude']).copy()
        df_matrix['lat_cluster'] = df_matrix['latitude'].round(5)
        df_matrix['lon_cluster'] = df_matrix['longitude'].round(5)
        df_matrix['lat_key'] = df_matrix['latitude'].round(2)
        df_matrix['lon_key'] = df_matrix['longitude'].round(2)

        # acq_date может быть смешанного типа (datetime/число/строка), поэтому сначала нормализуем.
        if 'acq_date' in df_matrix.columns:
            df_matrix['acq_date_safe'] = pd.to_datetime(df_matrix['acq_date'], errors='coerce')
        else:
            df_matrix['acq_date_safe'] = pd.NaT

        df_zones = df_matrix.groupby(
            ['Район', 'Населенный_пункт', 'lat_cluster', 'lon_cluster', 'lat_key', 'lon_key'],
            dropna=False
        ).agg({
            'acq_date_safe': 'max',
            'latitude': 'count'
        }).reset_index().rename(columns={
            'latitude': 'Кол_во_инцидентов',
            'acq_date_safe': 'acq_date'
        })

        df_fire = pd.read_excel('Риски_РТ_Для_QGIS.xlsx')
        df_zones = pd.merge(df_zones, df_fire[['Район', 'Индекс_Риска_R']], on='Район', how='left')
        df_zones.rename(columns={'Индекс_Риска_R': 'Индекс_Огня'}, inplace=True)

        df_flood = pd.read_excel('ТОП_РИСКИ_ПАВОДКОВ_ФИНАЛ.xlsx')
        df_flood['lat_key'] = pd.to_numeric(df_flood['latitude'], errors='coerce').round(2)
        df_flood['lon_key'] = pd.to_numeric(df_flood['longitude'], errors='coerce').round(2)
        df_flood_cluster = df_flood.groupby(['Район', 'Населенный_пункт', 'lat_key', 'lon_key'])[
            'Индекс_Риска_F'].max().reset_index()
        df_zones = pd.merge(df_zones, df_flood_cluster, on=['Район', 'Населенный_пункт', 'lat_key', 'lon_key'],
                            how='left')
        df_zones.rename(columns={'Индекс_Риска_F': 'Индекс_Воды'}, inplace=True)

        try:
            df_towers = pd.read_excel('Анализ_Связи_ПФО_ФИНАЛ (2).xlsx')
            df_towers = df_towers[df_towers['Регион'] == 'Республика Татарстан'].drop_duplicates(
                subset=['Населенный_пункт'])
            df_zones = pd.merge(df_zones,
                                df_towers[['Населенный_пункт', 'До_ближайшей_2G_вышки_км', 'До_ближайшей_4G_вышки_км']],
                                on='Населенный_пункт', how='left')
        except:
            pass
        df_zones['Старая_сирена'] = "НЕТ"
        # try:
        #     df_old = pd.read_excel('Справочник_ТСО_ФИНАЛ.xlsx')
        #     old_lats = np.radians(df_old['latitude'].values)
        #     old_lons = np.radians(df_old['longitude'].values)

        #     def check_siren_radius(lat, lon, radius_km=0.600):
        #         lat1, lon1 = np.radians(lat), np.radians(lon)
        #         a = np.sin((old_lats - lat1) / 2) ** 2 + np.cos(lat1) * np.cos(old_lats) * np.sin(
        #             (old_lons - lon1) / 2) ** 2
        #         return "ДА" if np.any(6371.0 * (2 * np.arctan2(np.sqrt(a), np.sqrt(1 - a))) <= radius_km) else "НЕТ"

        #     df_zones['Старая_сирена'] = df_zones.apply(lambda r: check_siren_radius(r['lat_cluster'], r['lon_cluster']),
        #                                                axis=1)
        # except:
        #     df_zones['Старая_сирена'] = "НЕТ"

        df_zones = df_zones.fillna({
            'Индекс_Огня': 0.0, 'Индекс_Воды': 0.0,
            'До_ближайшей_2G_вышки_км': 10.0, 'До_ближайшей_4G_вышки_км': 10.0, 'acq_date': 'Нет данных'
        })

        # Если несколько опасных зон находятся в радиусе 500 м, объединяем их в одну
        # усредненную опасную зону. Далее MILP-модель подбирает одну ТСО на эту группу.
        df_zones = aggregate_nearby_threat_zones(df_zones, radius_m=500)

        return df_zones
    except Exception as e:
        st.error(f"Ошибка загрузки файлов: {e}")
        return None
@st.cache_data
def load_old_tso():
    try:
        df_old = pd.read_excel('Справочник_ТСО_ФИНАЛ.xlsx')

        required_cols = ['latitude', 'longitude']
        for col in required_cols:
            if col not in df_old.columns:
                st.warning(f"В файле Справочник_ТСО_ФИНАЛ.xlsx нет колонки: {col}")
                return pd.DataFrame()

        df_old = df_old.dropna(subset=['latitude', 'longitude']).copy()

        df_old['latitude'] = pd.to_numeric(df_old['latitude'], errors='coerce')
        df_old['longitude'] = pd.to_numeric(df_old['longitude'], errors='coerce')
        df_old = df_old.dropna(subset=['latitude', 'longitude'])

        # Название населенного пункта
        if 'Населенный_пункт' in df_old.columns:
            df_old['Н.П.'] = df_old['Населенный_пункт']
        elif 'Район' in df_old.columns:
            df_old['Н.П.'] = df_old['Район']
        else:
            df_old['Н.П.'] = 'Нет данных'

        # Отдельная подсказка только для этих точек
        df_old['tooltip_html'] = (
            "<b>" + df_old['Н.П.'].astype(str) + "</b><br/>"
            "РАСЦО РТ"
        )

        return df_old

    except Exception as e:
        st.warning(f"Не удалось загрузить Справочник_ТСО_ФИНАЛ.xlsx: {e}")
        return pd.DataFrame()
@st.cache_data
def load_settlements():
    try:
        # sep=None сам определит разделитель: запятая, ; или табуляция
        df_np = pd.read_csv(
            'settlements_tatarstan_clean.csv',
            sep=None,
            engine='python',
            encoding='utf-8-sig'
        )

        required_cols = ['Район_OSM', 'Населенный_пункт_OSM', 'lat_np', 'lon_np', 'population_osm']
        for col in required_cols:
            if col not in df_np.columns:
                st.warning(f"В файле settlements_tatarstan_clean.csv нет колонки: {col}")
                return pd.DataFrame()

        df_np = df_np.copy()

        df_np['lat_np'] = pd.to_numeric(df_np['lat_np'], errors='coerce')
        df_np['lon_np'] = pd.to_numeric(df_np['lon_np'], errors='coerce')
        df_np['population_osm'] = pd.to_numeric(df_np['population_osm'], errors='coerce')

        df_np = df_np.dropna(subset=['lat_np', 'lon_np'])

        def calc_radius(row):
            name = str(row['Населенный_пункт_OSM']).strip().lower()
            pop = row['population_osm']

            # Специальные правила
            if name == 'казань':
                return 20000
            if name == 'набережные челны':
                return 7370

            # Если население не найдено
            if pd.isna(pop):
                return 700

            pop = float(pop)

            if pop <= 100:
                return 200
            elif pop <= 500:
                return 500
            elif pop <= 1500:
                return 700
            elif pop <= 5000:
                return 900
            elif pop <= 10000:
                return 1400
            elif pop <= 20000:
                return 2000
            elif pop <= 50000:
                return 4000
            elif pop <= 100000:
                return 4500
            elif pop <= 250000:
                return 6000
            elif pop <= 500000:
                return 7000
            elif pop <= 750000:
                return 8000
            elif pop <= 1000000:
                return 12000
            else:
                return 20000

        df_np['radius_m'] = df_np.apply(calc_radius, axis=1)
        # Площадь зоны населенного пункта, м²
        df_np['area_m2'] = math.pi * (df_np['radius_m'] ** 2)

        df_np['population_calc'] = df_np['population_osm'].fillna(0)
        
        df_np['people_per_m2'] = df_np['population_calc'] / df_np['area_m2']

        df_np['tooltip_html'] = (
            "<b>" + df_np['Населенный_пункт_OSM'].astype(str) + "</b><br/>"
            "Население: " + df_np['population_osm'].fillna('не найдено').astype(str) + "<br/>"
            "Радиус зоны: " + df_np['radius_m'].astype(str) + " м<br/>"
            "Плотность: " + df_np['people_per_m2'].round(6).astype(str) + " чел./м²"
        )

        return df_np

    except Exception as e:
        st.warning(f"Не удалось загрузить settlements_tatarstan_clean.csv: {e}")
        return pd.DataFrame()
@st.cache_data(show_spinner=False)
def add_population_to_threat_zones(df_zones, df_settlements):
    """
    Рассчитывает население, попадающее в каждую опасную зону.
    Опасная зона всегда имеет радиус 400 м.
    Население берется только из settlements_tatarstan_clean.csv.
    """

    if df_zones is None or df_zones.empty:
        return df_zones

    df_zones = df_zones.copy()
    df_zones['threat_radius_m'] = 400

    if df_settlements is None or df_settlements.empty:
        df_zones['Население'] = 0
        df_zones['Людей_в_опасной_зоне'] = 0
        df_zones['Площадь_пересечения_м2'] = 0
        return df_zones

    df_settlements = df_settlements.copy()

    # Приближенный перевод градусов в метры для Татарстана
    lat0 = 55.5
    meters_per_degree_lat = 111_320
    meters_per_degree_lon = 111_320 * math.cos(math.radians(lat0))

    def circle_intersection_area(r1, r2, d):
        """
        Площадь пересечения двух окружностей.
        r1 — радиус опасной зоны, м
        r2 — радиус зоны НП, м
        d — расстояние между центрами, м
        """

        if d >= r1 + r2:
            return 0.0

        if d <= abs(r1 - r2):
            return math.pi * min(r1, r2) ** 2

        part1 = r1 ** 2 * math.acos(
            (d ** 2 + r1 ** 2 - r2 ** 2) / (2 * d * r1)
        )

        part2 = r2 ** 2 * math.acos(
            (d ** 2 + r2 ** 2 - r1 ** 2) / (2 * d * r2)
        )

        part3 = 0.5 * math.sqrt(
            max(
                0,
                (-d + r1 + r2)
                * (d + r1 - r2)
                * (d - r1 + r2)
                * (d + r1 + r2)
            )
        )

        return part1 + part2 - part3

    settlements_valid = df_settlements.dropna(
        subset=['lat_np', 'lon_np', 'radius_m']
    ).copy()

    if settlements_valid.empty:
        df_zones['Население'] = 0
        df_zones['Людей_в_опасной_зоне'] = 0
        df_zones['Площадь_пересечения_м2'] = 0
        return df_zones

    # Координаты НП в метрах
    s_x = settlements_valid['lon_np'].astype(float).to_numpy() * meters_per_degree_lon
    s_y = settlements_valid['lat_np'].astype(float).to_numpy() * meters_per_degree_lat
    s_r = settlements_valid['radius_m'].astype(float).to_numpy()

    # Плотность населения, чел./м²
    s_density = settlements_valid['people_per_m2'].fillna(0).astype(float).to_numpy()

    threat_radius = 400

    population_results = []
    area_results = []

    for _, zone in df_zones.iterrows():
        if pd.isna(zone['lat_cluster']) or pd.isna(zone['lon_cluster']):
            population_results.append(0)
            area_results.append(0)
            continue

        z_x = float(zone['lon_cluster']) * meters_per_degree_lon
        z_y = float(zone['lat_cluster']) * meters_per_degree_lat

        dx = s_x - z_x
        dy = s_y - z_y
        distances = np.sqrt(dx * dx + dy * dy)

        # Берем только НП, окружности которых могут пересечься с угрозой 400 м
        candidate_mask = distances <= (s_r + threat_radius)

        if not np.any(candidate_mask):
            population_results.append(0)
            area_results.append(0)
            continue

        cand_distances = distances[candidate_mask]
        cand_radii = s_r[candidate_mask]
        cand_density = s_density[candidate_mask]

        total_people = 0.0
        total_area = 0.0

        for d, settlement_radius, density in zip(
            cand_distances,
            cand_radii,
            cand_density
        ):
            intersection_area = circle_intersection_area(
                threat_radius,
                settlement_radius,
                d
            )

            if intersection_area <= 0:
                continue

            total_area += intersection_area
            total_people += intersection_area * density

        population_results.append(int(round(total_people)))
        area_results.append(round(total_area, 2))

    df_zones['Население'] = population_results
    df_zones['Людей_в_опасной_зоне'] = population_results
    df_zones['Площадь_пересечения_м2'] = area_results

    return df_zones


def build_threat_tooltip(df_res):
    """Единый tooltip для всех опасных зон.

    Важно: рекомендованное ТСО выводится только в описании опасной зоны
    и в реестре. Отдельные новые точки ТСО на карте не создаются.
    """
    df_res = df_res.copy()

    for col in ['Людей_в_опасной_зоне', 'Площадь_пересечения_м2', 'Охват', 'Стоимость', 'Надежность']:
        if col not in df_res.columns:
            df_res[col] = 0
    if 'threat_radius_m' not in df_res.columns:
        df_res['threat_radius_m'] = 400
    if 'Кол_во_опасностей_в_группе' not in df_res.columns:
        df_res['Кол_во_опасностей_в_группе'] = 1
    if 'Кол_во_инцидентов' not in df_res.columns:
        df_res['Кол_во_инцидентов'] = 1

    # Совместимость со старой структурой: модель возвращает ТСО/Канал,
    # а пользовательский интерфейс показывает их именно как рекомендации.
    if 'Рекомендованное_ТСО' not in df_res.columns:
        df_res['Рекомендованное_ТСО'] = df_res.get('ТСО', 'ОТБРАКОВАНО')
    if 'Рекомендованный_канал' not in df_res.columns:
        df_res['Рекомендованный_канал'] = df_res.get('Канал', '-')

    rec_tso = df_res['Рекомендованное_ТСО'].fillna('ОТБРАКОВАНО').astype(str)
    rec_channel = df_res['Рекомендованный_канал'].fillna('-').astype(str)

    # Для зон, где модель не нашла допустимое оборудование, показываем понятное описание.
    rec_text = np.where(
        rec_tso.isin(['ОТБРАКОВАНО', 'nan', 'None', '']),
        'нет допустимого варианта по ограничениям',
        rec_tso + ' (' + rec_channel + ')'
    )
    df_res['Рекомендация_ТСО_текст'] = rec_text

    lat_text = pd.to_numeric(df_res['Широта'], errors='coerce').round(6).astype(str)
    lon_text = pd.to_numeric(df_res['Долгота'], errors='coerce').round(6).astype(str)

    df_res['tooltip_html'] = (
        "<b>Опасная зона</b><br/>"
        "<b>Н.П.:</b> " + df_res['Н.П.'].astype(str) + " (" + df_res['Район'].astype(str) + ")<br/>"
        "<b>Координаты:</b> " + lat_text + ", " + lon_text + "<br/>"
        "<b>Угроза:</b> " + df_res['Тип угрозы'].astype(str) + "<br/>"
        "<b>Радиус опасной зоны:</b> 400 м<br/>"
        "<b>Объединено зон в радиусе 500 м:</b> " + df_res['Кол_во_опасностей_в_группе'].fillna(1).round(0).astype(int).astype(str) + "<br/>"
        "<b>Количество инцидентов:</b> " + df_res['Кол_во_инцидентов'].fillna(1).round(0).astype(int).astype(str) + "<br/>"
        "<b>Людей в опасной зоне:</b> " + df_res['Людей_в_опасной_зоне'].fillna(0).round(0).astype(int).astype(str) + " чел.<br/>"
        "<b>Площадь пересечения с НП:</b> " + df_res['Площадь_пересечения_м2'].fillna(0).round(0).astype(int).astype(str) + " м²<br/>"
        "<b>Примечание:</b> ТСО рекомендуется для каждой зоны; если население не найдено, подбор выполняется по минимальной расчетной базе 50 чел.<br/>"
        "<b>Рекомендованное ТСО:</b> " + df_res['Рекомендация_ТСО_текст'].astype(str) + "<br/>"
        "<b>Стоимость:</b> " + df_res['Стоимость'].fillna(0).astype(str) + " у.е.<br/>"
        "<b>Расчетный охват:</b> " + df_res['Охват'].fillna(0).astype(str) + " чел.<br/>"
        "<b>Расчетная надежность:</b> " + df_res['Надежность'].fillna(0).astype(str)
    )
    return df_res

@st.cache_data(show_spinner=False)
def add_people_in_threat_zones_fast(df_threats, df_settlements):
    """
    Быстрый расчет людей в опасных зонах без shapely.
    Использует формулу площади пересечения двух окружностей.
    """

    if df_threats is None or df_threats.empty:
        return df_threats

    df_threats = df_threats.copy()

    if df_settlements is None or df_settlements.empty:
        df_threats['Людей_в_опасной_зоне'] = 0
        df_threats['Площадь_пересечения_м2'] = 0
        return df_threats

    df_settlements = df_settlements.copy()

    # Средняя широта Татарстана. Для расчета расстояний в метрах.
    lat0 = 55.5
    meters_per_degree_lat = 111_320
    meters_per_degree_lon = 111_320 * math.cos(math.radians(lat0))

    def circle_intersection_area(r1, r2, d):
        """
        Площадь пересечения двух окружностей.
        r1, r2, d — в метрах.
        """

        if d >= r1 + r2:
            return 0.0

        if d <= abs(r1 - r2):
            return math.pi * min(r1, r2) ** 2

        part1 = r1 ** 2 * math.acos((d ** 2 + r1 ** 2 - r2 ** 2) / (2 * d * r1))
        part2 = r2 ** 2 * math.acos((d ** 2 + r2 ** 2 - r1 ** 2) / (2 * d * r2))
        part3 = 0.5 * math.sqrt(
            max(
                0,
                (-d + r1 + r2)
                * (d + r1 - r2)
                * (d - r1 + r2)
                * (d + r1 + r2)
            )
        )

        return part1 + part2 - part3

    # Подготовка населенных пунктов в numpy-массивах
    settlements_valid = df_settlements.dropna(subset=['lat_np', 'lon_np', 'radius_m']).copy()

    if settlements_valid.empty:
        df_threats['Людей_в_опасной_зоне'] = 0
        df_threats['Площадь_пересечения_м2'] = 0
        return df_threats

    s_x = settlements_valid['lon_np'].astype(float).to_numpy() * meters_per_degree_lon
    s_y = settlements_valid['lat_np'].astype(float).to_numpy() * meters_per_degree_lat
    s_r = settlements_valid['radius_m'].astype(float).to_numpy()
    s_density = settlements_valid['people_per_m2'].fillna(0).astype(float).to_numpy()

    people_results = []
    area_results = []

    for _, threat in df_threats.iterrows():
        if pd.isna(threat['Широта']) or pd.isna(threat['Долгота']):
            people_results.append(0)
            area_results.append(0)
            continue

        threat_radius = float(threat.get('threat_radius_m', 400))

        t_x = float(threat['Долгота']) * meters_per_degree_lon
        t_y = float(threat['Широта']) * meters_per_degree_lat

        dx = s_x - t_x
        dy = s_y - t_y
        distances = np.sqrt(dx * dx + dy * dy)

        # Быстрый предварительный фильтр:
        # берем только те НП, окружности которых вообще могут пересечься с угрозой
        candidate_mask = distances <= (s_r + threat_radius)

        if not np.any(candidate_mask):
            people_results.append(0)
            area_results.append(0)
            continue

        cand_distances = distances[candidate_mask]
        cand_radii = s_r[candidate_mask]
        cand_density = s_density[candidate_mask]

        total_people = 0.0
        total_area = 0.0

        for d, settlement_radius, density in zip(cand_distances, cand_radii, cand_density):
            inter_area = circle_intersection_area(threat_radius, settlement_radius, d)

            if inter_area <= 0:
                continue

            total_area += inter_area
            total_people += inter_area * density

        people_results.append(int(round(total_people)))
        area_results.append(round(total_area, 2))

    df_threats['Людей_в_опасной_зоне'] = people_results
    df_threats['Площадь_пересечения_м2'] = area_results

    return df_threats
# --- 3. ЖЕСТКАЯ МАТЕМАТИЧЕСКАЯ МОДЕЛЬ ---
def run_optimization(df_zones, w_fire, w_flood, alpha, budget_large, budget_small, q_min, catalog_list):
    """
    ЖЕСТКАЯ МАТЕМАТИЧЕСКАЯ МОДЕЛЬ из пользовательского блока.

    Важно:
    - На вход подается уже усредненный слой опасностей: если несколько опасных зон
      находятся в радиусе 500 м, они заранее объединены в одну расчетную зону.
    - Для каждой такой усредненной зоны эта функция подбирает рекомендованное ТСО
      по исходной целевой функции PuLP.
    - Рекомендация рассчитывается для ВСЕХ опасных зон. Даже если население по
      пересечению с населенными пунктами равно 0, зона не отбраковывается:
      для выбора используется расчетное минимальное население pop_model=50,
      а в отчете сохраняется реальное население/охват.
    """
    catalog = catalog_list

    def calc_q_dyn(equip, r_f, r_w, d2, d4, pop):
        rel = equip["rel"]
        name = equip["name"]
        ch = equip["ch"]

        # ВАЖНО: ТСО нужно рекомендовать во всех опасных зонах.
        # Поэтому население меньше 50 не обнуляет Q, а используется только
        # как фактор выбора. Для правил подбора берем минимальную расчетную
        # базу 50 человек, но фактическое население/охват в отчете не меняем.
        pop_model = max(float(pop or 0), 50.0)

        boost = 0.0
        if name == "ТВ-системы" and ch == "Спутник" and pop_model >= 50000: boost = 10.0
        elif name == "ТВ-системы" and ch == "ТВ/радио" and 20000 <= pop_model < 50000: boost = 10.0
        elif name == "Речевые уст." and ch == "Проводной" and 10000 <= pop_model < 20000: boost = 10.0
        elif name == "Речевые уст." and ch == "IP" and 5000 <= pop_model < 10000: boost = 10.0
        elif name == "Речевые уст." and ch == "Сотовая" and 3000 <= pop_model < 5000: boost = 10.0
        elif name == "БАС (Дроны)" and ch == "Спутник" and r_f > 0.80 and d2 > 8.0: boost = 10.0
        elif name == "БАС (Дроны)" and ch == "Сотовая" and r_f > 0.80 and d2 <= 8.0: boost = 10.0
        elif name == "Мобильные комп." and ch == "Спутник" and r_w > 0.80 and d2 > 8.0 and pop_model > 500: boost = 10.0
        elif name == "Мобильные комп." and ch == "Сотовая" and r_w > 0.80 and d2 <= 8.0 and pop_model > 500: boost = 10.0

        if boost == 0.0:
            if "Моб. приложения" in name:
                if ch == "IP" and d4 <= 3.0: boost = 8.0
                elif ch == "Сотовая" and 3.0 < d4 <= 8.0: boost = 8.0
                elif ch == "Сотовая" and d4 > 8.0 and int(pop_model) % 3 == 0: boost = 8.0
            elif name == "SMS-оповещение" and ch == "Сотовая":
                if d4 > 8.0 and int(pop_model) % 3 == 1: boost = 8.0
            elif "Радиовещание" in name:
                if d4 > 8.0 and int(pop_model) % 3 == 2:
                    if ch == "Радио" and pop_model < 1000: boost = 8.0
                    elif ch == "ТВ/радио" and 1000 <= pop_model < 3000: boost = 8.0
                    elif ch == "Спутник" and pop_model >= 3000: boost = 8.0
            elif "Моб. приложения" in name and ch == "Сотовая":
                boost = 2.0

        return rel + boost

    prob = pulp.LpProblem("Final_Optimization", pulp.LpMinimize)
    vars_dict, obj_terms, init_risk = {}, [], 0

    for j, row in df_zones.iterrows():
        vars_dict[j] = {}
        R_base = (w_fire * row['Индекс_Огня']) + (w_flood * row['Индекс_Воды'])
        init_risk += R_base

        if row['Старая_сирена'] == "ДА":
            for k in range(len(catalog)):
                v = pulp.LpVariable(f"z_{j}_{k}", cat=pulp.LpBinary)
                vars_dict[j][k] = v
                prob += v == 0
            continue

        pop_real = float(row.get('Население', 0) or 0)
        # Для подбора рекомендации используем минимум 50 человек, чтобы зоны
        # без найденного населения не отбраковывались. Реальный охват ниже
        # по-прежнему считается от pop_real.
        pop_model = max(pop_real, 50.0)

        curr_b = budget_small if pop_model <= 500 else budget_large
        valid_keys = []

        for k, equip in enumerate(catalog):
            v = pulp.LpVariable(f"z_{j}_{k}", cat=pulp.LpBinary)
            vars_dict[j][k] = v
            Q = calc_q_dyn(equip, row['Индекс_Огня'], row['Индекс_Воды'], row['До_ближайшей_2G_вышки_км'],
                           row['До_ближайшей_4G_вышки_км'], pop_model)

            if equip["cost"] <= curr_b and Q >= q_min:
                valid_keys.append(k)
                # Оценочный охват для целевой функции: если фактическое население
                # равно 0, используем pop_model только для выбора подходящего ТСО.
                # В отчете фактический охват останется 0.
                O_model = min(equip["cov"], pop_model) * equip["k_act"]
                tf = max(0.1, (3600 - equip["time"]) / 3600.0)
                red = (Q * O_model * tf) / (pop_model * alpha + 1)
                obj_terms.append(-R_base * red * v)
            else:
                prob += v == 0

        prob += pulp.lpSum([vars_dict[j][k] for k in range(len(catalog))]) <= 1
        if valid_keys:
            prob += pulp.lpSum([vars_dict[j][k] for k in valid_keys]) == 1

    prob += pulp.lpSum(obj_terms)
    prob.solve(pulp.PULP_CBC_CMD(msg=0))

    report = []
    for j, row in df_zones.iterrows():
        th = "МУЛЬТИРИСК" if row['Индекс_Огня'] > 0.4 and row['Индекс_Воды'] > 0.4 else "ПОЖАР" if row['Индекс_Огня'] > row['Индекс_Воды'] else "ПАВОДОК"
        if th == "МУЛЬТИРИСК": c = [128, 0, 128, 200]
        elif th == "ПАВОДОК": c = [50, 100, 255, 200]
        else: c = [255, 50, 50, 200]
        
        pop_int = int(row['Население'])
        people_in_zone = int(row.get('Людей_в_опасной_зоне', pop_int))
        intersect_area = float(row.get('Площадь_пересечения_м2', 0))
        incident_count = int(row.get('Кол_во_инцидентов', 1))
        group_count = int(row.get('Кол_во_опасностей_в_группе', 1))

        base_report = {
            "Район": row['Район'],
            "Н.П.": row['Населенный_пункт'],
            "Широта": row['lat_cluster'],
            "Долгота": row['lon_cluster'],
            "Население": pop_int,
            "Население_для_подбора_ТСО": max(pop_int, 50),
            "Тип угрозы": th,
            "Людей_в_опасной_зоне": people_in_zone,
            "Площадь_пересечения_м2": intersect_area,
            "Кол_во_инцидентов": incident_count,
            "Кол_во_опасностей_в_группе": group_count,
            "threat_radius_m": 400,
            "color": c
        }

        if row['Старая_сирена'] == "ДА":
            item = dict(base_report)
            item.update({
                "ТСО": "ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ",
                "Канал": "Существующий",
                "Рекомендованное_ТСО": "ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ",
                "Рекомендованный_канал": "Существующий",
                "Стоимость": 0,
                "Охват": pop_int,
                "Надежность": 1.0
            })
            report.append(item)
            continue

        winner_found = False
        for k, equip in enumerate(catalog):
            if pulp.value(vars_dict[j][k]) is not None and pulp.value(vars_dict[j][k]) > 0.5:
                Q_report = equip["rel"] - (pop_int % 5) * 0.01 - (row['До_ближайшей_2G_вышки_км'] % 3) * 0.01
                Q_report = max(0.85, min(0.99, Q_report))

                O_final = int(min(equip["cov"], pop_int) * equip["k_act"])
                item = dict(base_report)
                item.update({
                    "ТСО": equip['name'],
                    "Канал": equip['ch'],
                    "Рекомендованное_ТСО": equip['name'],
                    "Рекомендованный_канал": equip['ch'],
                    "Стоимость": equip['cost'],
                    "Охват": O_final,
                    "Надежность": round(Q_report, 3)
                })
                report.append(item)
                winner_found = True

        if not winner_found:
            # Важно: зона все равно остается на карте, даже если население 0 или модель не нашла допустимую ТСО.
            # Цвет сохраняется по типу угрозы, чтобы опасная зона не пропадала визуально.
            item = dict(base_report)
            item.update({
                "ТСО": "ОТБРАКОВАНО",
                "Канал": "-",
                "Рекомендованное_ТСО": "ОТБРАКОВАНО",
                "Рекомендованный_канал": "-",
                "Стоимость": 0,
                "Охват": 0,
                "Надежность": 0.0
            })
            report.append(item)

    final_obj = init_risk + (pulp.value(prob.objective) or 0)
    return pd.DataFrame(report), init_risk, final_obj


# =============================================================================
# МЕХАНИЗМ ОБРАТНОЙ СВЯЗИ: ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =============================================================================

def _feedback_to_float(value, default=0.0, min_value=None, max_value=None):
    """Безопасно преобразует значение в число и ограничивает диапазон."""
    try:
        result = float(pd.to_numeric(value, errors='coerce'))
        if not np.isfinite(result):
            result = float(default)
    except Exception:
        result = float(default)

    if min_value is not None:
        result = max(float(min_value), result)
    if max_value is not None:
        result = min(float(max_value), result)
    return result


def _feedback_zone_key(district, settlement, latitude, longitude):
    """Формирует устойчивый ключ опасной зоны без изменения её координат."""
    lat = _feedback_to_float(latitude, 0.0)
    lon = _feedback_to_float(longitude, 0.0)
    return f"{str(district)}|{str(settlement)}|{lat:.6f}|{lon:.6f}"


def _feedback_normalize_weights(lambda_time, lambda_coverage, lambda_cost):
    """Нормализует три веса так, чтобы их сумма была строго равна единице."""
    values = np.array([
        _feedback_to_float(lambda_time, 0.35, 0.0, 1.0),
        _feedback_to_float(lambda_coverage, 0.40, 0.0, 1.0),
        _feedback_to_float(lambda_cost, 0.25, 0.0, 1.0)
    ], dtype=float)

    total = float(values.sum())
    if total <= 0:
        values = np.array([0.35, 0.40, 0.25], dtype=float)
    else:
        values = values / total

    # Два значения округляются, третье вычисляется как остаток. Это исключает
    # накопление ошибок округления и гарантирует сумму 1.000.
    lambda_time_norm = round(float(values[0]), 2)
    lambda_coverage_norm = round(float(values[1]), 2)
    lambda_cost_norm = round(max(0.0, 1.0 - lambda_time_norm - lambda_coverage_norm), 2)

    if lambda_cost_norm < 0:
        lambda_cost_norm = 0.0
        lambda_coverage_norm = round(max(0.0, 1.0 - lambda_time_norm), 2)

    total_after_round = lambda_time_norm + lambda_coverage_norm + lambda_cost_norm
    if abs(total_after_round - 1.0) > 1e-9:
        lambda_cost_norm = round(max(0.0, 1.0 - lambda_time_norm - lambda_coverage_norm), 2)

    return {
        'lambda_time': lambda_time_norm,
        'lambda_coverage': lambda_coverage_norm,
        'lambda_cost': lambda_cost_norm
    }


def _feedback_initialize_weight_state():
    """Создаёт сохраняемые значения весов и их предыдущие значения."""
    defaults = {
        'fb_lambda_time': 0.35,
        'fb_lambda_coverage': 0.40,
        'fb_lambda_cost': 0.25,
        '_fb_prev_lambda_time': 0.35,
        '_fb_prev_lambda_coverage': 0.40,
        '_fb_prev_lambda_cost': 0.25,
        'fb_initial_lambda_time': 0.35,
        'fb_initial_lambda_coverage': 0.40,
        'fb_initial_lambda_cost': 0.25
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def _feedback_redistribute_weights(changed_key):
    """Пропорционально перераспределяет остаток веса между двумя другими весами."""
    key_map = {
        'fb_lambda_time': '_fb_prev_lambda_time',
        'fb_lambda_coverage': '_fb_prev_lambda_coverage',
        'fb_lambda_cost': '_fb_prev_lambda_cost'
    }
    weight_keys = list(key_map.keys())

    changed_value = _feedback_to_float(st.session_state.get(changed_key, 0.0), 0.0, 0.0, 1.0)
    other_keys = [key for key in weight_keys if key != changed_key]
    old_other_values = [
        _feedback_to_float(st.session_state.get(key_map[key], 0.0), 0.0, 0.0, 1.0)
        for key in other_keys
    ]
    remainder = max(0.0, 1.0 - changed_value)
    old_other_total = float(sum(old_other_values))

    if old_other_total > 0:
        new_other_first = round(remainder * old_other_values[0] / old_other_total, 2)
        new_other_second = round(max(0.0, remainder - new_other_first), 2)
    else:
        new_other_first = round(remainder / 2.0, 2)
        new_other_second = round(max(0.0, remainder - new_other_first), 2)

    st.session_state[changed_key] = round(changed_value, 2)
    st.session_state[other_keys[0]] = new_other_first
    st.session_state[other_keys[1]] = new_other_second

    # Сохраняем новое состояние как исходное для следующего движения любого бегунка.
    for key, previous_key in key_map.items():
        st.session_state[previous_key] = _feedback_to_float(
            st.session_state.get(key, 0.0), 0.0, 0.0, 1.0
        )


def _feedback_get_weights_from_session():
    """Возвращает нормализованные веса, не изменяя состояния уже созданных виджетов."""
    return _feedback_normalize_weights(
        st.session_state.get('fb_lambda_time', 0.35),
        st.session_state.get('fb_lambda_coverage', 0.40),
        st.session_state.get('fb_lambda_cost', 0.25)
    )


def _feedback_catalog_lookup(catalog_list):
    """Создаёт справочник времени, стоимости, охвата и надежности по паре ТСО/канал."""
    lookup = {}
    for item in catalog_list or []:
        try:
            lookup[(str(item.get('name', '')), str(item.get('ch', '')))] = {
                'time': _feedback_to_float(item.get('time', 0.0), 0.0, 0.0),
                'cost': _feedback_to_float(item.get('cost', 0.0), 0.0, 0.0),
                'coverage': _feedback_to_float(item.get('cov', 0.0), 0.0, 0.0),
                'reliability': _feedback_to_float(item.get('rel', 0.0), 0.0, 0.0, 1.0),
                'activity': _feedback_to_float(item.get('k_act', 1.0), 1.0, 0.0, 1.0)
            }
        except Exception:
            continue
    return lookup


def create_feedback_input_table(df_res, catalog_list):
    """Создаёт исходную редактируемую таблицу фактических показателей по зонам."""
    display_columns = [
        'Район', 'Н.П.', 'Широта', 'Долгота', 'Тип угрозы',
        'Рекомендованное ТСО', 'Рекомендованный канал',
        'Плановое время', 'Фактическое время',
        'Плановый охват', 'Фактический охват',
        'Плановая стоимость', 'Фактическая стоимость',
        'Плановая надежность', 'Фактическая надежность ТСО',
        'Фактическая надежность канала'
    ]

    if df_res is None or df_res.empty:
        empty = pd.DataFrame(columns=display_columns)
        empty.index.name = 'Ключ зоны'
        return empty

    result = df_res.copy()
    for col, default in {
        'Район': 'Нет данных',
        'Н.П.': 'Нет данных',
        'Широта': 0.0,
        'Долгота': 0.0,
        'Тип угрозы': 'Нет данных',
        'Рекомендованное_ТСО': 'ОТБРАКОВАНО',
        'Рекомендованный_канал': '-',
        'ТСО': 'ОТБРАКОВАНО',
        'Канал': '-',
        'Охват': 0.0,
        'Стоимость': 0.0,
        'Надежность': 0.0
    }.items():
        if col not in result.columns:
            result[col] = default

    result['Рекомендованное ТСО'] = result['Рекомендованное_ТСО'].fillna(
        result['ТСО']
    ).fillna('ОТБРАКОВАНО').astype(str)
    result['Рекомендованный канал'] = result['Рекомендованный_канал'].fillna(
        result['Канал']
    ).fillna('-').astype(str)

    catalog_lookup = _feedback_catalog_lookup(catalog_list)

    def get_planned_time(row):
        tso = str(row.get('Рекомендованное ТСО', ''))
        channel = str(row.get('Рекомендованный канал', ''))
        if tso in ['ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ', 'ОТБРАКОВАНО']:
            return 0.0
        return _feedback_to_float(catalog_lookup.get((tso, channel), {}).get('time', 0.0), 0.0, 0.0)

    result['Плановое время'] = result.apply(get_planned_time, axis=1)
    result['Плановый охват'] = pd.to_numeric(result['Охват'], errors='coerce').fillna(0).clip(lower=0)
    result['Плановая стоимость'] = pd.to_numeric(result['Стоимость'], errors='coerce').fillna(0).clip(lower=0)
    result['Плановая надежность'] = pd.to_numeric(result['Надежность'], errors='coerce').fillna(0).clip(0, 1)

    result['Фактическое время'] = result['Плановое время']
    result['Фактический охват'] = result['Плановый охват']
    result['Фактическая стоимость'] = result['Плановая стоимость']
    result['Фактическая надежность ТСО'] = result['Плановая надежность']
    result['Фактическая надежность канала'] = result['Плановая надежность']

    zone_keys = result.apply(
        lambda row: _feedback_zone_key(
            row.get('Район', 'Нет данных'),
            row.get('Н.П.', 'Нет данных'),
            row.get('Широта', 0.0),
            row.get('Долгота', 0.0)
        ),
        axis=1
    )
    feedback_table = result[display_columns].copy()
    feedback_table.index = zone_keys
    feedback_table.index.name = 'Ключ зоны'
    return feedback_table



def calculate_adapted_feedback_weights(feedback_table, selected_weights, weight_adaptation_coefficient):
    """Обновляет веса по фактическому отклонению показателей и нормализует их.

    Коэффициент адаптации весов определяет, насколько сильно фактическое ухудшение
    показателя повышает приоритет соответствующего критерия при повторной оптимизации.
    """
    base = _feedback_normalize_weights(
        selected_weights.get('lambda_time', 0.35),
        selected_weights.get('lambda_coverage', 0.40),
        selected_weights.get('lambda_cost', 0.25)
    )
    if feedback_table is None or feedback_table.empty:
        return base

    frame = feedback_table.copy()

    def total(column):
        if column not in frame.columns:
            return 0.0
        return float(pd.to_numeric(frame[column], errors='coerce').fillna(0).sum())

    plan_time = total('Плановое время')
    fact_time = total('Фактическое время')
    plan_coverage = total('Плановый охват')
    fact_coverage = total('Фактический охват')
    plan_cost = total('Плановая стоимость')
    fact_cost = total('Фактическая стоимость')
    coefficient = _feedback_to_float(weight_adaptation_coefficient, 0.50, 0.0, 1.0)

    time_deviation = max(0.0, fact_time / max(plan_time, 1.0) - 1.0)
    coverage_deviation = max(0.0, 1.0 - fact_coverage / max(plan_coverage, 1.0))
    cost_deviation = max(0.0, fact_cost / max(plan_cost, 1.0) - 1.0)

    return _feedback_normalize_weights(
        base['lambda_time'] * (1.0 + coefficient * time_deviation),
        base['lambda_coverage'] * (1.0 + coefficient * coverage_deviation),
        base['lambda_cost'] * (1.0 + coefficient * cost_deviation)
    )


def calculate_feedback_zone_parameters(feedback_table, adaptation_parameters):
    """Рассчитывает допустимые обновленные параметры по данным обратной связи."""
    if feedback_table is None or feedback_table.empty:
        return pd.DataFrame()

    result = feedback_table.copy().reset_index().rename(columns={'index': 'Ключ зоны'})
    if 'Ключ зоны' not in result.columns:
        result['Ключ зоны'] = result.apply(
            lambda row: _feedback_zone_key(
                row.get('Район', 'Нет данных'),
                row.get('Н.П.', 'Нет данных'),
                row.get('Широта', 0.0),
                row.get('Долгота', 0.0)
            ),
            axis=1
        )

    non_negative_columns = [
        'Плановое время', 'Фактическое время',
        'Плановый охват', 'Фактический охват',
        'Плановая стоимость', 'Фактическая стоимость'
    ]
    reliability_columns = [
        'Плановая надежность',
        'Фактическая надежность ТСО',
        'Фактическая надежность канала'
    ]
    for col in non_negative_columns:
        if col not in result.columns:
            result[col] = 0.0
        result[col] = pd.to_numeric(result[col], errors='coerce').fillna(0).clip(lower=0)
    for col in reliability_columns:
        if col not in result.columns:
            result[col] = 0.0
        result[col] = pd.to_numeric(result[col], errors='coerce').fillna(0).clip(0, 1)

    k_time = _feedback_to_float(adaptation_parameters.get('k_time', 0.50), 0.50, 0.0, 1.0)
    k_coverage = _feedback_to_float(adaptation_parameters.get('k_coverage', 0.50), 0.50, 0.0, 1.0)
    k_cost = _feedback_to_float(adaptation_parameters.get('k_cost', 0.50), 0.50, 0.0, 1.0)
    k_tso_reliability = _feedback_to_float(adaptation_parameters.get('k_tso_reliability', 0.50), 0.50, 0.0, 1.0)
    k_channel_reliability = _feedback_to_float(adaptation_parameters.get('k_channel_reliability', 0.50), 0.50, 0.0, 1.0)

    # Формулы реализуют сглаженное обновление: p_new = p_plan + k * (p_fact - p_plan).
    # Риски и их веса в этих формулах намеренно не участвуют.
    result['Время_для_оптимизации'] = (
        result['Плановое время'] + k_time * (result['Фактическое время'] - result['Плановое время'])
    ).clip(lower=0)
    result['Охват_для_оптимизации'] = (
        result['Плановый охват'] + k_coverage * (result['Фактический охват'] - result['Плановый охват'])
    ).clip(lower=0)
    result['Стоимость_для_оптимизации'] = (
        result['Плановая стоимость'] + k_cost * (result['Фактическая стоимость'] - result['Плановая стоимость'])
    ).clip(lower=0)
    result['Надежность_ТСО_для_оптимизации'] = (
        result['Плановая надежность'] + k_tso_reliability * (
            result['Фактическая надежность ТСО'] - result['Плановая надежность']
        )
    ).clip(0, 1)
    result['Надежность_канала_для_оптимизации'] = (
        result['Плановая надежность'] + k_channel_reliability * (
            result['Фактическая надежность канала'] - result['Плановая надежность']
        )
    ).clip(0, 1)
    result['Надежность_для_оптимизации'] = np.minimum(
        result['Надежность_ТСО_для_оптимизации'],
        result['Надежность_канала_для_оптимизации']
    ).clip(0, 1)

    return result


def _feedback_ratio(updated_value, planned_value, default=1.0):
    """Безопасно рассчитывает коэффициент изменения показателя."""
    updated = _feedback_to_float(updated_value, default, 0.0)
    planned = _feedback_to_float(planned_value, 0.0, 0.0)
    if planned <= 0:
        return float(default)
    return max(0.0, updated / planned)


def run_feedback_optimization(
    df_zones,
    feedback_parameters_df,
    w_fire,
    w_flood,
    alpha,
    budget_large,
    budget_small,
    q_min,
    catalog_list,
    objective_weights
):
    """Выполняет повторную оптимизацию с обновленными эксплуатационными параметрами.

    Имеет отдельную функцию, чтобы базовая оптимизация не менялась. Комплексный риск
    всегда повторно берётся из исходных `Индекс_Огня`, `Индекс_Воды`, `w_fire`, `w_flood`.
    """
    if df_zones is None or df_zones.empty:
        return pd.DataFrame(), 0.0, 0.0

    catalog = [dict(item) for item in (catalog_list or [])]
    if not catalog:
        return pd.DataFrame(), 0.0, 0.0

    weights = _feedback_normalize_weights(
        objective_weights.get('lambda_time', 0.35),
        objective_weights.get('lambda_coverage', 0.40),
        objective_weights.get('lambda_cost', 0.25)
    )
    lambda_time = weights['lambda_time']
    lambda_coverage = weights['lambda_coverage']
    lambda_cost = weights['lambda_cost']

    feedback_lookup = {}
    if feedback_parameters_df is not None and not feedback_parameters_df.empty:
        source = feedback_parameters_df.copy()
        if 'Ключ зоны' not in source.columns:
            source['Ключ зоны'] = source.apply(
                lambda row: _feedback_zone_key(
                    row.get('Район', 'Нет данных'),
                    row.get('Н.П.', 'Нет данных'),
                    row.get('Широта', 0.0),
                    row.get('Долгота', 0.0)
                ),
                axis=1
            )
        feedback_lookup = source.drop_duplicates('Ключ зоны').set_index('Ключ зоны').to_dict('index')

    max_time = max([_feedback_to_float(item.get('time', 0), 0.0, 0.0) for item in catalog] + [1.0])
    prob = pulp.LpProblem('Feedback_Optimization', pulp.LpMinimize)
    vars_dict = {}
    candidate_values = {}
    obj_terms = []
    init_risk = 0.0

    def calc_q_dyn(equip, r_f, r_w, d2, d4, pop):
        rel = _feedback_to_float(equip.get('rel', 0.0), 0.0, 0.0, 1.0)
        name = str(equip.get('name', ''))
        channel = str(equip.get('ch', ''))
        pop_model = max(_feedback_to_float(pop, 50.0, 0.0), 50.0)
        boost = 0.0

        if name == 'ТВ-системы' and channel == 'Спутник' and pop_model >= 50000:
            boost = 10.0
        elif name == 'ТВ-системы' and channel == 'ТВ/радио' and 20000 <= pop_model < 50000:
            boost = 10.0
        elif name == 'Речевые уст.' and channel == 'Проводной' and 10000 <= pop_model < 20000:
            boost = 10.0
        elif name == 'Речевые уст.' and channel == 'IP' and 5000 <= pop_model < 10000:
            boost = 10.0
        elif name == 'Речевые уст.' and channel == 'Сотовая' and 3000 <= pop_model < 5000:
            boost = 10.0
        elif name == 'БАС (Дроны)' and channel == 'Спутник' and r_f > 0.80 and d2 > 8.0:
            boost = 10.0
        elif name == 'БАС (Дроны)' and channel == 'Сотовая' and r_f > 0.80 and d2 <= 8.0:
            boost = 10.0
        elif name == 'Мобильные комп.' and channel == 'Спутник' and r_w > 0.80 and d2 > 8.0 and pop_model > 500:
            boost = 10.0
        elif name == 'Мобильные комп.' and channel == 'Сотовая' and r_w > 0.80 and d2 <= 8.0 and pop_model > 500:
            boost = 10.0

        if boost == 0.0:
            if 'Моб. приложения' in name:
                if channel == 'IP' and d4 <= 3.0:
                    boost = 8.0
                elif channel == 'Сотовая' and 3.0 < d4 <= 8.0:
                    boost = 8.0
                elif channel == 'Сотовая' and d4 > 8.0 and int(pop_model) % 3 == 0:
                    boost = 8.0
            elif name == 'SMS-оповещение' and channel == 'Сотовая' and d4 > 8.0 and int(pop_model) % 3 == 1:
                boost = 8.0
            elif 'Радиовещание' in name and d4 > 8.0 and int(pop_model) % 3 == 2:
                if channel == 'Радио' and pop_model < 1000:
                    boost = 8.0
                elif channel == 'ТВ/радио' and 1000 <= pop_model < 3000:
                    boost = 8.0
                elif channel == 'Спутник' and pop_model >= 3000:
                    boost = 8.0

        return rel + boost

    for j, row in df_zones.iterrows():
        vars_dict[j] = {}
        risk_fire = _feedback_to_float(row.get('Индекс_Огня', 0.0), 0.0)
        risk_flood = _feedback_to_float(row.get('Индекс_Воды', 0.0), 0.0)
        # ВАЖНО: данный расчёт полностью сохраняет исходную формулу комплексного риска.
        risk_base = w_fire * risk_fire + w_flood * risk_flood
        init_risk += risk_base

        zone_key = _feedback_zone_key(
            row.get('Район', 'Нет данных'),
            row.get('Населенный_пункт', 'Нет данных'),
            row.get('lat_cluster', 0.0),
            row.get('lon_cluster', 0.0)
        )
        feedback_row = feedback_lookup.get(zone_key, {})

        planned_time = _feedback_to_float(feedback_row.get('Плановое время', 0.0), 0.0, 0.0)
        updated_time = _feedback_to_float(feedback_row.get('Время_для_оптимизации', planned_time), planned_time, 0.0)
        planned_coverage = _feedback_to_float(feedback_row.get('Плановый охват', 0.0), 0.0, 0.0)
        updated_coverage = _feedback_to_float(feedback_row.get('Охват_для_оптимизации', planned_coverage), planned_coverage, 0.0)
        planned_cost = _feedback_to_float(feedback_row.get('Плановая стоимость', 0.0), 0.0, 0.0)
        updated_cost = _feedback_to_float(feedback_row.get('Стоимость_для_оптимизации', planned_cost), planned_cost, 0.0)
        planned_reliability = _feedback_to_float(feedback_row.get('Плановая надежность', 0.0), 0.0, 0.0, 1.0)
        updated_tso_reliability = _feedback_to_float(
            feedback_row.get('Надежность_ТСО_для_оптимизации', planned_reliability), planned_reliability, 0.0, 1.0
        )
        updated_channel_reliability = _feedback_to_float(
            feedback_row.get('Надежность_канала_для_оптимизации', planned_reliability), planned_reliability, 0.0, 1.0
        )

        time_factor = _feedback_ratio(updated_time, planned_time, 1.0)
        coverage_factor = _feedback_ratio(updated_coverage, planned_coverage, 1.0)
        cost_factor = _feedback_ratio(updated_cost, planned_cost, 1.0)
        tso_reliability_factor = _feedback_ratio(updated_tso_reliability, planned_reliability, 1.0)
        channel_reliability_factor = _feedback_ratio(updated_channel_reliability, planned_reliability, 1.0)

        if str(row.get('Старая_сирена', 'НЕТ')) == 'ДА':
            for k in range(len(catalog)):
                variable = pulp.LpVariable(f'feedback_z_{j}_{k}', cat=pulp.LpBinary)
                vars_dict[j][k] = variable
                prob += variable == 0
            continue

        population_real = _feedback_to_float(row.get('Население', 0.0), 0.0, 0.0)
        population_model = max(population_real, 50.0)
        local_budget = budget_small if population_model <= 500 else budget_large
        valid_keys = []

        for k, equip in enumerate(catalog):
            variable = pulp.LpVariable(f'feedback_z_{j}_{k}', cat=pulp.LpBinary)
            vars_dict[j][k] = variable

            raw_dynamic_reliability = calc_q_dyn(
                equip,
                risk_fire,
                risk_flood,
                _feedback_to_float(row.get('До_ближайшей_2G_вышки_км', 10.0), 10.0, 0.0),
                _feedback_to_float(row.get('До_ближайшей_4G_вышки_км', 10.0), 10.0, 0.0),
                population_model
            )
            candidate_time = _feedback_to_float(equip.get('time', 0.0), 0.0, 0.0) * time_factor
            candidate_cost = _feedback_to_float(equip.get('cost', 0.0), 0.0, 0.0) * cost_factor
            candidate_coverage = min(
                _feedback_to_float(equip.get('cov', 0.0), 0.0, 0.0)
                * _feedback_to_float(equip.get('k_act', 1.0), 1.0, 0.0, 1.0)
                * coverage_factor,
                population_model
            )
            candidate_reliability = _feedback_to_float(equip.get('rel', 0.0), 0.0, 0.0, 1.0)
            candidate_reliability *= tso_reliability_factor * channel_reliability_factor
            candidate_reliability = _feedback_to_float(candidate_reliability, 0.0, 0.0, 1.0)

            candidate_values[(j, k)] = {
                'time': candidate_time,
                'cost': candidate_cost,
                'coverage': candidate_coverage,
                'reliability': candidate_reliability
            }

            # Техническая допустимость опирается на исходный динамический критерий
            # и актуализированную фактическую надежность. Риск при этом не корректируется.
            if (
                candidate_cost <= _feedback_to_float(local_budget, 0.0, 0.0)
                and raw_dynamic_reliability >= _feedback_to_float(q_min, 0.60, 0.0, 1.0)
                and candidate_reliability >= _feedback_to_float(q_min, 0.60, 0.0, 1.0)
            ):
                valid_keys.append(k)
                time_norm = candidate_time / max(max_time, 1.0)
                coverage_norm = candidate_coverage / max(population_model, 1.0)
                cost_norm = candidate_cost / max(_feedback_to_float(local_budget, 1.0, 1.0), 1.0)
                # Формула повторной оптимизации: время и стоимость минимизируются,
                # охват и надежность максимизируются; комплексный риск сохранён исходным.
                objective_value = (
                    lambda_time * time_norm
                    - lambda_coverage * coverage_norm
                    + lambda_cost * cost_norm
                    - 0.10 * candidate_reliability
                )
                obj_terms.append(risk_base * objective_value * variable)
            else:
                prob += variable == 0

        prob += pulp.lpSum([vars_dict[j][k] for k in range(len(catalog))]) <= 1
        if valid_keys:
            prob += pulp.lpSum([vars_dict[j][k] for k in valid_keys]) == 1

    prob += pulp.lpSum(obj_terms)
    prob.solve(pulp.PULP_CBC_CMD(msg=0))

    report = []
    for j, row in df_zones.iterrows():
        risk_fire = _feedback_to_float(row.get('Индекс_Огня', 0.0), 0.0)
        risk_flood = _feedback_to_float(row.get('Индекс_Воды', 0.0), 0.0)
        threat_type = 'МУЛЬТИРИСК' if risk_fire > 0.4 and risk_flood > 0.4 else 'ПОЖАР' if risk_fire > risk_flood else 'ПАВОДОК'
        if threat_type == 'МУЛЬТИРИСК':
            color = [128, 0, 128, 200]
        elif threat_type == 'ПАВОДОК':
            color = [50, 100, 255, 200]
        else:
            color = [255, 50, 50, 200]

        population_real = int(_feedback_to_float(row.get('Население', 0.0), 0.0, 0.0))
        base_report = {
            'Район': row.get('Район', 'Нет данных'),
            'Н.П.': row.get('Населенный_пункт', 'Нет данных'),
            'Широта': _feedback_to_float(row.get('lat_cluster', 0.0), 0.0),
            'Долгота': _feedback_to_float(row.get('lon_cluster', 0.0), 0.0),
            'Население': population_real,
            'Население_для_подбора_ТСО': max(population_real, 50),
            'Тип угрозы': threat_type,
            'Людей_в_опасной_зоне': int(_feedback_to_float(row.get('Людей_в_опасной_зоне', population_real), population_real, 0.0)),
            'Площадь_пересечения_м2': _feedback_to_float(row.get('Площадь_пересечения_м2', 0.0), 0.0, 0.0),
            'Кол_во_инцидентов': int(_feedback_to_float(row.get('Кол_во_инцидентов', 1), 1.0, 1.0)),
            'Кол_во_опасностей_в_группе': int(_feedback_to_float(row.get('Кол_во_опасностей_в_группе', 1), 1.0, 1.0)),
            'threat_radius_m': 400,
            'color': color,
            'Индекс_Огня': risk_fire,
            'Индекс_Воды': risk_flood
        }

        if str(row.get('Старая_сирена', 'НЕТ')) == 'ДА':
            item = dict(base_report)
            item.update({
                'ТСО': 'ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ',
                'Канал': 'Существующий',
                'Рекомендованное_ТСО': 'ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ',
                'Рекомендованный_канал': 'Существующий',
                'Стоимость': 0.0,
                'Охват': population_real,
                'Надежность': 1.0,
                'Время оповещения': 0.0
            })
            report.append(item)
            continue

        winner_found = False
        for k, equip in enumerate(catalog):
            if pulp.value(vars_dict[j][k]) is not None and pulp.value(vars_dict[j][k]) > 0.5:
                values = candidate_values.get((j, k), {})
                item = dict(base_report)
                item.update({
                    'ТСО': str(equip.get('name', 'ОТБРАКОВАНО')),
                    'Канал': str(equip.get('ch', '-')),
                    'Рекомендованное_ТСО': str(equip.get('name', 'ОТБРАКОВАНО')),
                    'Рекомендованный_канал': str(equip.get('ch', '-')),
                    'Стоимость': round(_feedback_to_float(values.get('cost', 0.0), 0.0, 0.0), 2),
                    'Охват': int(round(min(_feedback_to_float(values.get('coverage', 0.0), 0.0, 0.0), population_real))),
                    'Надежность': round(_feedback_to_float(values.get('reliability', 0.0), 0.0, 0.0, 1.0), 3),
                    'Время оповещения': round(_feedback_to_float(values.get('time', 0.0), 0.0, 0.0), 2)
                })
                report.append(item)
                winner_found = True

        if not winner_found:
            item = dict(base_report)
            item.update({
                'ТСО': 'ОТБРАКОВАНО',
                'Канал': '-',
                'Рекомендованное_ТСО': 'ОТБРАКОВАНО',
                'Рекомендованный_канал': '-',
                'Стоимость': 0.0,
                'Охват': 0,
                'Надежность': 0.0,
                'Время оповещения': 0.0
            })
            report.append(item)

    final_objective = float(pulp.value(prob.objective) or 0.0)
    return pd.DataFrame(report), init_risk, final_objective


def calculate_feedback_comparison_table(df_before_feedback, df_after_feedback, feedback_parameters_df):
    """Сравнивает исходные и адаптированные решения для каждой опасной зоны."""
    columns = [
        'Район', 'Н.П.', 'Широта', 'Долгота', 'Тип угрозы',
        'ТСО до', 'Канал до', 'ТСО после', 'Канал после',
        'Время до', 'Время после', 'Охват до', 'Охват после',
        'Стоимость до', 'Стоимость после', 'Надежность до', 'Надежность после',
        'Изменение ТСО', 'Изменение канала', 'Изменение конфигурации', 'Вывод по зоне'
    ]
    if df_before_feedback is None or df_before_feedback.empty:
        return pd.DataFrame(columns=columns)

    before = df_before_feedback.copy().reset_index(drop=True)
    after = df_after_feedback.copy() if df_after_feedback is not None else pd.DataFrame()
    feedback = feedback_parameters_df.copy() if feedback_parameters_df is not None else pd.DataFrame()

    before['Ключ зоны'] = before.apply(
        lambda row: _feedback_zone_key(
            row.get('Район', 'Нет данных'), row.get('Н.П.', 'Нет данных'),
            row.get('Широта', 0.0), row.get('Долгота', 0.0)
        ), axis=1
    )
    if not after.empty:
        after['Ключ зоны'] = after.apply(
            lambda row: _feedback_zone_key(
                row.get('Район', 'Нет данных'), row.get('Н.П.', 'Нет данных'),
                row.get('Широта', 0.0), row.get('Долгота', 0.0)
            ), axis=1
        )
        after_lookup = after.drop_duplicates('Ключ зоны').set_index('Ключ зоны').to_dict('index')
    else:
        after_lookup = {}

    if not feedback.empty:
        if 'Ключ зоны' not in feedback.columns:
            feedback['Ключ зоны'] = feedback.apply(
                lambda row: _feedback_zone_key(
                    row.get('Район', 'Нет данных'), row.get('Н.П.', 'Нет данных'),
                    row.get('Широта', 0.0), row.get('Долгота', 0.0)
                ), axis=1
            )
        feedback_lookup = feedback.drop_duplicates('Ключ зоны').set_index('Ключ зоны').to_dict('index')
    else:
        feedback_lookup = {}

    rows = []
    for _, row in before.iterrows():
        key = row['Ключ зоны']
        after_row = after_lookup.get(key, {})
        fact_row = feedback_lookup.get(key, {})

        tso_before = str(row.get('Рекомендованное ТСО', 'ОТБРАКОВАНО'))
        channel_before = str(row.get('Рекомендованный канал', '-'))
        tso_after = str(after_row.get('Рекомендованное_ТСО', after_row.get('ТСО', 'ОТБРАКОВАНО')))
        channel_after = str(after_row.get('Рекомендованный_канал', after_row.get('Канал', '-')))
        tso_changed = 'Да' if tso_before != tso_after else 'Нет'
        channel_changed = 'Да' if channel_before != channel_after else 'Нет'
        configuration_changed = 'Да' if tso_changed == 'Да' or channel_changed == 'Да' else 'Нет'

        plan_time = _feedback_to_float(row.get('Плановое время', 0.0), 0.0, 0.0)
        plan_coverage = _feedback_to_float(row.get('Плановый охват', 0.0), 0.0, 0.0)
        plan_cost = _feedback_to_float(row.get('Плановая стоимость', 0.0), 0.0, 0.0)
        plan_reliability = _feedback_to_float(row.get('Плановая надежность', 0.0), 0.0, 0.0, 1.0)
        fact_time = _feedback_to_float(fact_row.get('Фактическое время', plan_time), plan_time, 0.0)
        fact_coverage = _feedback_to_float(fact_row.get('Фактический охват', plan_coverage), plan_coverage, 0.0)
        fact_cost = _feedback_to_float(fact_row.get('Фактическая стоимость', plan_cost), plan_cost, 0.0)
        fact_tso_reliability = _feedback_to_float(
            fact_row.get('Фактическая надежность ТСО', plan_reliability), plan_reliability, 0.0, 1.0
        )
        fact_channel_reliability = _feedback_to_float(
            fact_row.get('Фактическая надежность канала', plan_reliability), plan_reliability, 0.0, 1.0
        )
        local_problem = (
            fact_time > plan_time + 1e-9
            or fact_coverage < plan_coverage - 1e-9
            or min(fact_tso_reliability, fact_channel_reliability) < plan_reliability - 1e-9
            or fact_cost > plan_cost + 1e-9
        )

        if configuration_changed == 'Да':
            conclusion = 'Конфигурация изменена повторной оптимизацией.'
        elif local_problem:
            conclusion = 'Выявлено локальное отклонение фактических показателей; требуется контроль.'
        else:
            conclusion = 'Конфигурация сохранена; существенных локальных отклонений нет.'

        rows.append({
            'Район': row.get('Район', 'Нет данных'),
            'Н.П.': row.get('Н.П.', 'Нет данных'),
            'Широта': _feedback_to_float(row.get('Широта', 0.0), 0.0),
            'Долгота': _feedback_to_float(row.get('Долгота', 0.0), 0.0),
            'Тип угрозы': row.get('Тип угрозы', 'Нет данных'),
            'ТСО до': tso_before,
            'Канал до': channel_before,
            'ТСО после': tso_after,
            'Канал после': channel_after,
            'Время до': round(plan_time, 2),
            'Время после': round(_feedback_to_float(after_row.get('Время оповещения', plan_time), plan_time, 0.0), 2),
            'Охват до': round(plan_coverage, 2),
            'Охват после': round(_feedback_to_float(after_row.get('Охват', 0.0), 0.0, 0.0), 2),
            'Стоимость до': round(plan_cost, 2),
            'Стоимость после': round(_feedback_to_float(after_row.get('Стоимость', 0.0), 0.0, 0.0), 2),
            'Надежность до': round(plan_reliability, 3),
            'Надежность после': round(_feedback_to_float(after_row.get('Надежность', 0.0), 0.0, 0.0, 1.0), 3),
            'Изменение ТСО': tso_changed,
            'Изменение канала': channel_changed,
            'Изменение конфигурации': configuration_changed,
            'Вывод по зоне': conclusion
        })

    return pd.DataFrame(rows, columns=columns)



def calculate_comparable_feedback_objective(
    solution_df,
    df_zones,
    objective_weights,
    budget_large,
    budget_small,
    time_column,
    coverage_column,
    cost_column,
    reliability_column,
    common_time_normalizer
):
    """Рассчитывает сопоставимое значение целевой функции для решений до и после.

    Использует одинаковые нормировочные коэффициенты и исключительно исходные индексы
    риска. Поэтому значения можно корректно сравнивать между собой.
    """
    if solution_df is None or solution_df.empty or df_zones is None or df_zones.empty:
        return 0.0

    weights = _feedback_normalize_weights(
        objective_weights.get('lambda_time', 0.35),
        objective_weights.get('lambda_coverage', 0.40),
        objective_weights.get('lambda_cost', 0.25)
    )

    risk_lookup = {}
    for _, zone in df_zones.iterrows():
        key = _feedback_zone_key(
            zone.get('Район', 'Нет данных'),
            zone.get('Населенный_пункт', 'Нет данных'),
            zone.get('lat_cluster', 0.0),
            zone.get('lon_cluster', 0.0)
        )
        risk_lookup[key] = {
            'risk_fire': _feedback_to_float(zone.get('Индекс_Огня', 0.0), 0.0),
            'risk_flood': _feedback_to_float(zone.get('Индекс_Воды', 0.0), 0.0),
            'population': _feedback_to_float(zone.get('Население', 0.0), 0.0, 0.0)
        }

    total_objective = 0.0
    normalizer = max(_feedback_to_float(common_time_normalizer, 1.0, 1.0), 1.0)
    for _, row in solution_df.iterrows():
        key = _feedback_zone_key(
            row.get('Район', 'Нет данных'), row.get('Н.П.', 'Нет данных'),
            row.get('Широта', 0.0), row.get('Долгота', 0.0)
        )
        source = risk_lookup.get(key, {})
        population = max(_feedback_to_float(source.get('population', 0.0), 0.0, 0.0), 50.0)
        local_budget = budget_small if population <= 500 else budget_large
        risk = (
            _feedback_to_float(st.session_state.get('w_fire_feedback_reference', 0.6), 0.6, 0.0, 1.0)
            * _feedback_to_float(source.get('risk_fire', 0.0), 0.0)
            + _feedback_to_float(st.session_state.get('w_flood_feedback_reference', 0.4), 0.4, 0.0, 1.0)
            * _feedback_to_float(source.get('risk_flood', 0.0), 0.0)
        )
        time_norm = _feedback_to_float(row.get(time_column, 0.0), 0.0, 0.0) / normalizer
        coverage_norm = _feedback_to_float(row.get(coverage_column, 0.0), 0.0, 0.0) / population
        cost_norm = _feedback_to_float(row.get(cost_column, 0.0), 0.0, 0.0) / max(_feedback_to_float(local_budget, 1.0, 1.0), 1.0)
        reliability = _feedback_to_float(row.get(reliability_column, 0.0), 0.0, 0.0, 1.0)
        total_objective += risk * (
            weights['lambda_time'] * time_norm
            - weights['lambda_coverage'] * coverage_norm
            + weights['lambda_cost'] * cost_norm
            - 0.10 * reliability
        )
    return float(total_objective)


def calculate_feedback_metrics(df_before_feedback, df_after_feedback, objective_before, objective_after):
    """Подготавливает общие показатели до и после обратной связи."""
    before = df_before_feedback.copy() if df_before_feedback is not None else pd.DataFrame()
    after = df_after_feedback.copy() if df_after_feedback is not None else pd.DataFrame()

    def safe_sum(frame, column):
        if frame.empty or column not in frame.columns:
            return 0.0
        return float(pd.to_numeric(frame[column], errors='coerce').fillna(0).sum())

    def safe_mean(frame, column):
        if frame.empty or column not in frame.columns:
            return 0.0
        values = pd.to_numeric(frame[column], errors='coerce').dropna()
        return float(values.mean()) if not values.empty else 0.0

    cost_before = safe_sum(before, 'Плановая стоимость')
    coverage_before = safe_sum(before, 'Плановый охват')
    time_before = safe_mean(before, 'Плановое время')
    reliability_before = safe_mean(before, 'Плановая надежность')
    cost_after = safe_sum(after, 'Стоимость')
    coverage_after = safe_sum(after, 'Охват')
    time_after = safe_mean(after, 'Время оповещения')
    reliability_after = safe_mean(after, 'Надежность')

    changed_tso = 0
    changed_channel = 0
    changed_configuration = 0
    comparison = st.session_state.get('feedback_comparison_df', pd.DataFrame())
    if isinstance(comparison, pd.DataFrame) and not comparison.empty:
        changed_tso = int((comparison['Изменение ТСО'] == 'Да').sum())
        changed_channel = int((comparison['Изменение канала'] == 'Да').sum())
        changed_configuration = int((comparison['Изменение конфигурации'] == 'Да').sum())

    rows = [
        {'Показатель': 'Суммарная стоимость', 'До': cost_before, 'После': cost_after, 'Изменение': cost_after - cost_before},
        {'Показатель': 'Суммарный охват', 'До': coverage_before, 'После': coverage_after, 'Изменение': coverage_after - coverage_before},
        {'Показатель': 'Среднее время оповещения', 'До': time_before, 'После': time_after, 'Изменение': time_after - time_before},
        {'Показатель': 'Средняя надежность', 'До': reliability_before, 'После': reliability_after, 'Изменение': reliability_after - reliability_before},
        {'Показатель': 'Сопоставимое значение целевой функции', 'До': objective_before, 'После': objective_after, 'Изменение': objective_after - objective_before},
        {'Показатель': 'Зон с изменённым ТСО', 'До': 0, 'После': changed_tso, 'Изменение': changed_tso},
        {'Показатель': 'Зон с изменённым каналом', 'До': 0, 'После': changed_channel, 'Изменение': changed_channel},
        {'Показатель': 'Зон с изменённой конфигурацией', 'До': 0, 'После': changed_configuration, 'Изменение': changed_configuration}
    ]
    return pd.DataFrame(rows)


def calculate_feedback_parameter_changes(feedback_parameters_df, weights):
    """Формирует таблицу изменения управляемых параметров обратной связи."""
    if feedback_parameters_df is None or feedback_parameters_df.empty:
        return pd.DataFrame(columns=['Параметр', 'Значение до', 'Значение после', 'Изменение', 'Комментарий'])

    frame = feedback_parameters_df.copy()

    def mean_value(column):
        if column not in frame.columns:
            return 0.0
        values = pd.to_numeric(frame[column], errors='coerce').dropna()
        return float(values.mean()) if not values.empty else 0.0

    initial_weights = {
        'lambda_time': _feedback_to_float(st.session_state.get('fb_initial_lambda_time', 0.35), 0.35, 0.0, 1.0),
        'lambda_coverage': _feedback_to_float(st.session_state.get('fb_initial_lambda_coverage', 0.40), 0.40, 0.0, 1.0),
        'lambda_cost': _feedback_to_float(st.session_state.get('fb_initial_lambda_cost', 0.25), 0.25, 0.0, 1.0)
    }

    rows = [
        {'Параметр': 'Среднее время оповещения', 'Значение до': mean_value('Плановое время'), 'Значение после': mean_value('Время_для_оптимизации'), 'Комментарий': 'Сглаженное обновление по фактическому времени.'},
        {'Параметр': 'Средний охват населения', 'Значение до': mean_value('Плановый охват'), 'Значение после': mean_value('Охват_для_оптимизации'), 'Комментарий': 'Сглаженное обновление по фактическому охвату.'},
        {'Параметр': 'Средняя стоимость', 'Значение до': mean_value('Плановая стоимость'), 'Значение после': mean_value('Стоимость_для_оптимизации'), 'Комментарий': 'Сглаженное обновление по фактической стоимости.'},
        {'Параметр': 'Средняя надежность ТСО', 'Значение до': mean_value('Плановая надежность'), 'Значение после': mean_value('Надежность_ТСО_для_оптимизации'), 'Комментарий': 'Сглаженное обновление по фактической надежности ТСО.'},
        {'Параметр': 'Средняя надежность канала', 'Значение до': mean_value('Плановая надежность'), 'Значение после': mean_value('Надежность_канала_для_оптимизации'), 'Комментарий': 'Сглаженное обновление по фактической надежности канала.'},
        {'Параметр': 'λ1 — вес времени', 'Значение до': initial_weights['lambda_time'], 'Значение после': weights['lambda_time'], 'Комментарий': 'Нормализованный вес времени в повторной оптимизации.'},
        {'Параметр': 'λ2 — вес охвата', 'Значение до': initial_weights['lambda_coverage'], 'Значение после': weights['lambda_coverage'], 'Комментарий': 'Нормализованный вес охвата в повторной оптимизации.'},
        {'Параметр': 'λ3 — вес стоимости', 'Значение до': initial_weights['lambda_cost'], 'Значение после': weights['lambda_cost'], 'Комментарий': 'Нормализованный вес стоимости в повторной оптимизации.'}
    ]
    result = pd.DataFrame(rows)
    result['Изменение'] = result['Значение после'] - result['Значение до']
    return result[['Параметр', 'Значение до', 'Значение после', 'Изменение', 'Комментарий']]


def determine_feedback_decision(metrics_df, comparison_df, chi, minimum_efficiency_gain):
    """Принимает решение о полной замене, частичной корректировке или сохранении."""
    if metrics_df is None or metrics_df.empty:
        return 'Сохранение текущей конфигурации', 'Недостаточно данных для оценки эффекта обратной связи.'

    metric_map = metrics_df.set_index('Показатель').to_dict('index')
    time_before = _feedback_to_float(metric_map.get('Среднее время оповещения', {}).get('До', 0.0), 0.0, 0.0)
    time_after = _feedback_to_float(metric_map.get('Среднее время оповещения', {}).get('После', 0.0), 0.0, 0.0)
    coverage_before = _feedback_to_float(metric_map.get('Суммарный охват', {}).get('До', 0.0), 0.0, 0.0)
    coverage_after = _feedback_to_float(metric_map.get('Суммарный охват', {}).get('После', 0.0), 0.0, 0.0)
    cost_before = _feedback_to_float(metric_map.get('Суммарная стоимость', {}).get('До', 0.0), 0.0, 0.0)
    cost_after = _feedback_to_float(metric_map.get('Суммарная стоимость', {}).get('После', 0.0), 0.0, 0.0)
    objective_delta = _feedback_to_float(metric_map.get('Сопоставимое значение целевой функции', {}).get('Изменение', 0.0), 0.0)

    efficiency_time = (time_before - time_after) / max(time_before, 1.0)
    efficiency_coverage = (coverage_after - coverage_before) / max(coverage_before, 1.0)
    efficiency_cost = (cost_before - cost_after) / max(cost_before, 1.0)
    minimum_gain = _feedback_to_float(minimum_efficiency_gain, 0.0, 0.0, 1.0)
    significance_threshold = _feedback_to_float(chi, 0.0, 0.0, 1.0)

    local_problems = 0
    if isinstance(comparison_df, pd.DataFrame) and not comparison_df.empty:
        local_problems = int(comparison_df['Вывод по зоне'].astype(str).str.contains('локальное отклонение').sum())

    all_metrics_improved = (
        efficiency_time >= minimum_gain
        and efficiency_coverage >= minimum_gain
        and efficiency_cost >= minimum_gain
    )

    if objective_delta < -significance_threshold and all_metrics_improved:
        return (
            'Полная замена конфигурации',
            'Новое решение существенно лучше: снижение сопоставимой целевой функции превышает порог χ, '
            'а показатели времени, охвата и стоимости достигли минимального прироста эффективности.'
        )
    if local_problems > 0 or (objective_delta < 0 and abs(objective_delta) <= significance_threshold):
        return (
            'Частичная корректировка конфигурации',
            'Обнаружены локальные проблемы либо улучшение недостаточно велико для полной замены. '
            'Рекомендуется корректировать конфигурации только в проблемных зонах.'
        )
    return (
        'Сохранение текущей конфигурации',
        'Повторная оптимизация не дала значимого улучшения либо новое решение не лучше исходного. '
        'Текущую конфигурацию рекомендуется сохранить.'
    )


def build_feedback_excel(df_before, df_after, comparison_df, parameter_changes_df, metrics_df):
    """Создаёт Excel-файл с результатами обратной связи в памяти."""
    output = BytesIO()
    with pd.ExcelWriter(output) as writer:
        (df_before if isinstance(df_before, pd.DataFrame) else pd.DataFrame()).to_excel(
            writer, sheet_name='До обратной связи', index=False
        )
        (df_after if isinstance(df_after, pd.DataFrame) else pd.DataFrame()).to_excel(
            writer, sheet_name='После обратной связи', index=False
        )
        (comparison_df if isinstance(comparison_df, pd.DataFrame) else pd.DataFrame()).to_excel(
            writer, sheet_name='Сравнение по зонам', index=False
        )
        (parameter_changes_df if isinstance(parameter_changes_df, pd.DataFrame) else pd.DataFrame()).to_excel(
            writer, sheet_name='Параметры обратной связи', index=False
        )
        (metrics_df if isinstance(metrics_df, pd.DataFrame) else pd.DataFrame()).to_excel(
            writer, sheet_name='Итоговые показатели', index=False
        )
    output.seek(0)
    return output.getvalue()


def reset_feedback_session_state():
    """Удаляет результаты предыдущего цикла обратной связи при новом базовом расчете."""
    keys_to_remove = [
        'feedback_enabled', 'feedback_input_df', 'feedback_parameters_df',
        'df_before_feedback', 'df_feedback_res', 'feedback_comparison_df',
        'feedback_parameter_changes_df', 'feedback_metrics_df', 'feedback_decision',
        'feedback_decision_comment', 'feedback_r_in', 'feedback_r_out',
        'fb_lambda_time', 'fb_lambda_coverage', 'fb_lambda_cost',
        '_fb_prev_lambda_time', '_fb_prev_lambda_coverage', '_fb_prev_lambda_cost',
        'fb_initial_lambda_time', 'fb_initial_lambda_coverage', 'fb_initial_lambda_cost',
        'fb_k_time', 'fb_k_coverage', 'fb_k_cost', 'fb_k_tso_reliability',
        'fb_k_channel_reliability', 'fb_k_weights', 'fb_chi', 'fb_min_efficiency',
        'feedback_data_editor', 'w_fire_feedback_reference', 'w_flood_feedback_reference',
        'feedback_weights_after_adaptation'
    ]
    for key in keys_to_remove:
        if key in st.session_state:
            del st.session_state[key]


# --- ИНТЕРФЕЙС STREAMLIT ---
boundary_data = get_tatarstan_geojson()

settlements_points = load_settlements()
old_tso_points = load_old_tso()

data_result_raw = load_data()

if data_result_raw is not None:
    data_result = add_population_to_threat_zones(
        data_result_raw,
        settlements_points
    )
else:
    data_result = None

if data_result is not None:
    # ===== ЛЕВАЯ ПАНЕЛЬ (SIDEBAR) =====
    st.sidebar.header("Глобальные системные константы")
    
    w_flood = st.sidebar.slider("Вес риска наводнений", 0.0, 1.0, 0.4, 0.1)
    w_fire = round(1.0 - w_flood, 1)
    st.sidebar.markdown(f"""
        Вес риска пожаров: 
        ### {w_fire}""")

    alpha = 0.9
    q_min = st.sidebar.slider("Порог надежности ТСО (Q_min)", 0.1, 0.9, 0.60, 0.05)

    st.sidebar.markdown("**Локальный бюджетный лимит:**")
    b_max_large = st.sidebar.number_input("Бюджет (Население > 500 чел)", 1000, 20000, 3000, 500)
    b_max_small = st.sidebar.number_input("Бюджет (Население ≤ 500 чел)", 50, 5000, 500, 50)

    st.sidebar.markdown("---")
    st.sidebar.header("Константы ТСО")
    st.sidebar.markdown("*(Выберите систему из списка, чтобы изменить 5 параметров)*")

    tso_names = [f"{item['name']} ({item['ch']})" for item in st.session_state.catalog]
    selected_tso_name = st.sidebar.selectbox("Выберите ТСО", tso_names)
    selected_idx = tso_names.index(selected_tso_name)
    current_item = st.session_state.catalog[selected_idx]

    new_cost = st.sidebar.number_input("Cost (Стоимость, у.е.)", min_value=10, max_value=20000,
                                       value=int(current_item['cost']), step=50)
    new_cov = st.sidebar.number_input("Cov (Тех. охват, чел)", min_value=100, max_value=50000,
                                      value=int(current_item['cov']), step=500)
    new_rel = st.sidebar.slider("Rel_base (Базовая надежность)", 0.50, 0.99, float(current_item['rel']), 0.01)
    new_time = st.sidebar.number_input("Time (Время срабатывания, сек)", min_value=1, max_value=3600, value=int(current_item['time']), step=1)
    #new_k_act = st.sidebar.slider("K_act (Коэф. вовлеченности)", 0.10, 1.00, float(current_item['k_act']), 0.05)

    st.session_state.catalog[selected_idx].update({
        'cost': new_cost,
        'cov': new_cov,
        'rel': new_rel,
        'time': new_time
        #'k_act': new_k_act
    })

    #if st.sidebar.button("Сбросить данные к заводским", use_container_width=True):
    #   st.session_state.catalog = [dict(item) for item in DEFAULT_CATALOG]
    #    st.rerun()

    # ===== ОСНОВНАЯ ОБЛАСТЬ ЭКРАНА =====
    st.subheader("Технологические ограничения структуры каналов связи")
    
    st.subheader("Справочник оборудования")
    df_display_catalog = pd.DataFrame(st.session_state.catalog)
    df_display_catalog.columns = ['Система', 'Канал связи', 'Cost (Стоимость)', 'Cov (Охват)', 'Rel_base (Надежность)','Time (Время)', 'K_act (Вовлеченность)']
    df_display_catalog = df_display_catalog.drop(columns=['K_act (Вовлеченность)'])
    st.dataframe(df_display_catalog, use_container_width=True, hide_index=True)

    st.markdown("---")

    if st.button("ЗАПУСТИТЬ ОПТИМИЗАЦИЮ", type="primary"):
        with st.spinner("Ожидайте. Идет расчет глобального оптимума..."):
            df_res, r_in, r_out = run_optimization(data_result, w_fire, w_flood, alpha, b_max_large, b_max_small, q_min,
                                                   st.session_state.catalog)

        df_res['Вероятность_ошибки'] = 1.0 - df_res['Надежность']
        summary_table = df_res.groupby(['ТСО', 'Канал']).agg(
            Количество=('ТСО', 'count'), Общая_стоимость=('Стоимость', 'sum'),
            Общий_охват=('Охват', 'sum'), Ср_надежность=('Надежность', 'mean'), Ср_ошибка=('Вероятность_ошибки', 'mean')
        ).reset_index()
        summary_table['Система_и_Канал'] = summary_table['ТСО'] + " (" + summary_table['Канал'] + ")"

        st.session_state['opt_run'] = True
        st.session_state['summary_table'] = summary_table
        st.session_state['df_res'] = df_res
        st.session_state['r_in'] = r_in
        st.session_state['r_out'] = r_out
        reset_feedback_session_state()

    if st.session_state.get('opt_run', False):
        df_res = st.session_state['df_res']
        summary_table = st.session_state['summary_table']
        r_in = st.session_state['r_in']
        r_out = st.session_state['r_out']

        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Всего кластеров", len(df_res))
        col2.metric("Зон с рекомендацией ТСО", len(df_res[~df_res['Рекомендованное_ТСО'].isin(['ОБОРУДОВАНО СТАРОЙ СИРЕНОЙ', 'ОТБРАКОВАНО'])]))
        col3.metric("Общий бюджет (у.е.)", f"{df_res['Стоимость'].sum():,}")
        #col4.metric("Общий риск", f"{r_in:.2f}")
        
        # Tooltip и служебные поля для карты. Пересечения уже рассчитаны ДО оптимизации
        # в data_result и перенесены в df_res внутри run_optimization.
        df_res = build_threat_tooltip(df_res)

        # КАРТА
        layers = []

        # 1. Границы Татарстана — нижний слой
        if boundary_data:
            layers.append(
                pdk.Layer(
                    "GeoJsonLayer",
                    boundary_data,
                    opacity=0.3,
                    stroked=True,
                    filled=True,
                    get_fill_color=[100, 150, 200, 20],
                    get_line_color=[100, 100, 100, 150],
                    line_width_min_pixels=1,
                    pickable=False
                )
            )

        # 2. Зоны населенных пунктов по численности населения.
        # Они нужны как фон, поэтому pickable=False: это ускоряет карту и не мешает наведению на опасности/ТСО.
        if settlements_points is not None and not settlements_points.empty:
            layers.append(
                pdk.Layer(
                    "ScatterplotLayer",
                    data=settlements_points,
                    get_position=["lon_np", "lat_np"],
                    get_radius="radius_m",
                    get_fill_color=[0, 120, 255, 25],
                    get_line_color=[0, 80, 200, 90],
                    stroked=True,
                    filled=True,
                    line_width_min_pixels=1,
                    pickable=False
                )
            )

            # Центры населенных пунктов — маленькие синие точки с tooltip.
            layers.append(
                pdk.Layer(
                    "ScatterplotLayer",
                    data=settlements_points,
                    get_position=["lon_np", "lat_np"],
                    get_radius=80,
                    get_fill_color=[0, 40, 180, 220],
                    pickable=True,
                    filled=True
                )
            )

        # 3. Все зоны опасностей. Никакой фильтрации по населению/пересечению нет.
        # Даже если людей в зоне 0, круг остается на карте и окрашен по типу угрозы.
        df_threats_map = df_res.copy()
        df_threats_map['threat_radius_m'] = 400

        if not df_threats_map.empty:
            layers.append(
                pdk.Layer(
                    "ScatterplotLayer",
                    data=df_threats_map,
                    get_position=["Долгота", "Широта"],
                    get_color="color",
                    get_radius="threat_radius_m",
                    pickable=True,
                    filled=True
                )
            )

        # 4. Все существующие ТСО из Справочник_ТСО_ФИНАЛ.xlsx.
        # Это отдельный слой отображения; он не исключает опасности из расчета.
        if old_tso_points is not None and not old_tso_points.empty:
            layers.append(
                pdk.Layer(
                    "ScatterplotLayer",
                    data=old_tso_points,
                    get_position=["longitude", "latitude"],
                    get_radius=300,
                    get_fill_color=[50, 200, 50, 35],
                    get_line_color=[20, 160, 20, 255],
                    stroked=True,
                    filled=True,
                    line_width_min_pixels=2,
                    pickable=True
                )
            )

        st.pydeck_chart(
            pdk.Deck(
                layers=layers,
                initial_view_state=pdk.ViewState(
                    latitude=df_res['Широта'].mean(),
                    longitude=df_res['Долгота'].mean(),
                    zoom=6
                ),
                tooltip={
                    "html": "{tooltip_html}"
                }
            ),
            use_container_width=True,
            height=900
        )

        st.subheader("Реестр кластеров")
        st.dataframe(df_res, use_container_width=True)
        csv = df_res.to_csv(index=False, sep=';').encode('utf-8-sig')

        st.download_button(
            label="Скачать итоговый реестр (CSV)",
            data=csv,
            file_name="TSO_Optimization_Final.csv",
            mime="text/csv; charset=utf-8-sig"
        )

        # ==========================================
        # БЛОК РАСЧЕТА ОБРАТНОЙ СВЯЗИ
        # ==========================================
        st.markdown('---')
        st.subheader('Расчет обратной связи')
        st.caption(
            'Повторная оптимизация использует исходные риски без их пересчёта. '
            'Изменяются только фактические эксплуатационные показатели и веса целевой функции.'
        )

        feedback_enabled = st.checkbox(
            'Включить расчет обратной связи',
            key='feedback_enabled'
        )

        if feedback_enabled:
            _feedback_initialize_weight_state()

            if 'feedback_input_df' not in st.session_state:
                st.session_state['feedback_input_df'] = create_feedback_input_table(
                    st.session_state.get('df_res', pd.DataFrame()),
                    st.session_state.catalog
                )

            with st.expander('Параметры обратной связи', expanded=True):
                st.markdown('**Весовые коэффициенты целевой функции**')
                weight_col_1, weight_col_2, weight_col_3 = st.columns(3)
                with weight_col_1:
                    st.slider(
                        'λ1 — вес времени оповещения',
                        min_value=0.0,
                        max_value=1.0,
                        step=0.01,
                        key='fb_lambda_time',
                        on_change=_feedback_redistribute_weights,
                        args=('fb_lambda_time',),
                        help='При увеличении веса времени два остальных веса автоматически уменьшаются пропорционально.'
                    )
                with weight_col_2:
                    st.slider(
                        'λ2 — вес охвата населения',
                        min_value=0.0,
                        max_value=1.0,
                        step=0.01,
                        key='fb_lambda_coverage',
                        on_change=_feedback_redistribute_weights,
                        args=('fb_lambda_coverage',),
                        help='При увеличении веса охвата два остальных веса автоматически уменьшаются пропорционально.'
                    )
                with weight_col_3:
                    st.slider(
                        'λ3 — вес стоимости',
                        min_value=0.0,
                        max_value=1.0,
                        step=0.01,
                        key='fb_lambda_cost',
                        on_change=_feedback_redistribute_weights,
                        args=('fb_lambda_cost',),
                        help='При увеличении веса стоимости два остальных веса автоматически уменьшаются пропорционально.'
                    )

                feedback_weights = _feedback_get_weights_from_session()
                st.write(
                    f"Сумма весов: "
                    f"{feedback_weights['lambda_time'] + feedback_weights['lambda_coverage'] + feedback_weights['lambda_cost']:.3f}"
                )

                st.markdown('**Коэффициенты адаптации и критерии принятия решения**')
                adaptation_col_1, adaptation_col_2 = st.columns(2)
                with adaptation_col_1:
                    st.slider(
                        'Коэффициент адаптации времени', 0.0, 1.0, 0.50, 0.01,
                        key='fb_k_time',
                        help='Определяет, насколько фактическое время влияет на обновлённое время для повторной оптимизации.'
                    )
                    st.slider(
                        'Коэффициент адаптации охвата', 0.0, 1.0, 0.50, 0.01,
                        key='fb_k_coverage',
                        help='Определяет, насколько фактический охват влияет на обновлённый охват.'
                    )
                    st.slider(
                        'Коэффициент адаптации стоимости', 0.0, 1.0, 0.50, 0.01,
                        key='fb_k_cost',
                        help='Определяет, насколько фактическая стоимость влияет на обновлённую стоимость.'
                    )
                    st.slider(
                        'Коэффициент адаптации надежности ТСО', 0.0, 1.0, 0.50, 0.01,
                        key='fb_k_tso_reliability',
                        help='Определяет влияние фактической надежности технического средства на повторный расчёт.'
                    )
                with adaptation_col_2:
                    st.slider(
                        'Коэффициент адаптации надежности канала связи', 0.0, 1.0, 0.50, 0.01,
                        key='fb_k_channel_reliability',
                        help='Определяет влияние фактической надежности канала связи на повторный расчёт.'
                    )
                    st.slider(
                        'Коэффициент адаптации весов', 0.0, 1.0, 0.50, 0.01,
                        key='fb_k_weights',
                        help='Фиксируется в параметрах цикла обратной связи и используется для экспертной настройки весов.'
                    )
                    st.slider(
                        'Порог значимости χ', 0.0, 1.0, 0.05, 0.01,
                        key='fb_chi',
                        help='Минимальное значимое изменение сопоставимого значения целевой функции.'
                    )
                    st.slider(
                        'Минимальный прирост эффективности', 0.0, 1.0, 0.05, 0.01,
                        key='fb_min_efficiency',
                        help='Минимальная доля улучшения времени, охвата и стоимости для признания полной замены успешной.'
                    )

            st.markdown('**Фактические показатели по опасным зонам**')
            st.caption('Редактируются только фактические значения. Плановые и идентификационные поля заблокированы.')
            identifier_columns = [
                'Район', 'Н.П.', 'Широта', 'Долгота', 'Тип угрозы',
                'Рекомендованное ТСО', 'Рекомендованный канал',
                'Плановое время', 'Плановый охват',
                'Плановая стоимость', 'Плановая надежность'
            ]
            feedback_editor_df = st.data_editor(
                st.session_state['feedback_input_df'],
                key='feedback_data_editor',
                use_container_width=True,
                hide_index=True,
                num_rows='fixed',
                disabled=identifier_columns,
                column_config={
                    'Фактическое время': st.column_config.NumberColumn('Фактическое время', min_value=0.0, step=1.0),
                    'Фактический охват': st.column_config.NumberColumn('Фактический охват', min_value=0.0, step=1.0),
                    'Фактическая стоимость': st.column_config.NumberColumn('Фактическая стоимость', min_value=0.0, step=1.0),
                    'Фактическая надежность ТСО': st.column_config.NumberColumn('Фактическая надежность ТСО', min_value=0.0, max_value=1.0, step=0.01),
                    'Фактическая надежность канала': st.column_config.NumberColumn('Фактическая надежность канала', min_value=0.0, max_value=1.0, step=0.01)
                }
            )
            st.session_state['feedback_input_df'] = feedback_editor_df.copy()

            if st.button('РАССЧИТАТЬ ОБРАТНУЮ СВЯЗЬ', type='primary', use_container_width=True):
                with st.spinner('Выполняется адаптация параметров и повторная оптимизация...'):
                    adaptation_parameters = {
                        'k_time': st.session_state.get('fb_k_time', 0.50),
                        'k_coverage': st.session_state.get('fb_k_coverage', 0.50),
                        'k_cost': st.session_state.get('fb_k_cost', 0.50),
                        'k_tso_reliability': st.session_state.get('fb_k_tso_reliability', 0.50),
                        'k_channel_reliability': st.session_state.get('fb_k_channel_reliability', 0.50),
                        'k_weights': st.session_state.get('fb_k_weights', 0.50),
                        'chi': st.session_state.get('fb_chi', 0.05),
                        'minimum_efficiency_gain': st.session_state.get('fb_min_efficiency', 0.05)
                    }
                    df_before_feedback = feedback_editor_df.copy().reset_index(drop=True)
                    adapted_feedback_weights = calculate_adapted_feedback_weights(
                        feedback_editor_df,
                        feedback_weights,
                        adaptation_parameters['k_weights']
                    )
                    feedback_parameters_df = calculate_feedback_zone_parameters(
                        feedback_editor_df,
                        adaptation_parameters
                    )
                    st.session_state['w_fire_feedback_reference'] = w_fire
                    st.session_state['w_flood_feedback_reference'] = w_flood
                    df_feedback_res, feedback_r_in, feedback_r_out = run_feedback_optimization(
                        data_result.copy(),
                        feedback_parameters_df,
                        w_fire,
                        w_flood,
                        alpha,
                        b_max_large,
                        b_max_small,
                        q_min,
                        st.session_state.catalog,
                        adapted_feedback_weights
                    )
                    feedback_comparison_df = calculate_feedback_comparison_table(
                        df_before_feedback,
                        df_feedback_res,
                        feedback_parameters_df
                    )
                    common_time_normalizer = max(
                        _feedback_to_float(df_before_feedback.get('Плановое время', pd.Series([0])).max(), 0.0, 0.0),
                        _feedback_to_float(df_feedback_res.get('Время оповещения', pd.Series([0])).max(), 0.0, 0.0),
                        1.0
                    )
                    objective_before = calculate_comparable_feedback_objective(
                        df_before_feedback,
                        data_result,
                        adapted_feedback_weights,
                        b_max_large,
                        b_max_small,
                        'Плановое время',
                        'Плановый охват',
                        'Плановая стоимость',
                        'Плановая надежность',
                        common_time_normalizer
                    )
                    objective_after = calculate_comparable_feedback_objective(
                        df_feedback_res,
                        data_result,
                        adapted_feedback_weights,
                        b_max_large,
                        b_max_small,
                        'Время оповещения',
                        'Охват',
                        'Стоимость',
                        'Надежность',
                        common_time_normalizer
                    )
                    st.session_state['df_before_feedback'] = df_before_feedback
                    st.session_state['feedback_parameters_df'] = feedback_parameters_df
                    st.session_state['df_feedback_res'] = df_feedback_res
                    st.session_state['feedback_comparison_df'] = feedback_comparison_df
                    st.session_state['feedback_r_in'] = feedback_r_in
                    st.session_state['feedback_r_out'] = feedback_r_out
                    st.session_state['feedback_weights_after_adaptation'] = adapted_feedback_weights
                    st.session_state['feedback_parameter_changes_df'] = calculate_feedback_parameter_changes(
                        feedback_parameters_df,
                        adapted_feedback_weights
                    )
                    st.session_state['feedback_metrics_df'] = calculate_feedback_metrics(
                        df_before_feedback,
                        df_feedback_res,
                        objective_before,
                        objective_after
                    )
                    decision, decision_comment = determine_feedback_decision(
                        st.session_state['feedback_metrics_df'],
                        feedback_comparison_df,
                        adaptation_parameters['chi'],
                        adaptation_parameters['minimum_efficiency_gain']
                    )
                    st.session_state['feedback_decision'] = decision
                    st.session_state['feedback_decision_comment'] = decision_comment

                st.success('Расчет обратной связи завершен. Базовый df_res не изменялся.')

            if 'df_feedback_res' in st.session_state:
                feedback_metrics_df = st.session_state.get('feedback_metrics_df', pd.DataFrame())
                feedback_parameter_changes_df = st.session_state.get('feedback_parameter_changes_df', pd.DataFrame())
                feedback_comparison_df = st.session_state.get('feedback_comparison_df', pd.DataFrame())
                df_before_feedback = st.session_state.get('df_before_feedback', pd.DataFrame())
                df_feedback_res = st.session_state.get('df_feedback_res', pd.DataFrame())

                st.markdown('### Результаты обратной связи')
                if not feedback_metrics_df.empty:
                    metric_lookup = feedback_metrics_df.set_index('Показатель').to_dict('index')
                    metric_col_1, metric_col_2, metric_col_3, metric_col_4 = st.columns(4)
                    metric_col_1.metric(
                        'Суммарная стоимость',
                        f"{_feedback_to_float(metric_lookup.get('Суммарная стоимость', {}).get('После', 0.0), 0.0):,.2f}",
                        f"{_feedback_to_float(metric_lookup.get('Суммарная стоимость', {}).get('Изменение', 0.0), 0.0):+,.2f}"
                    )
                    metric_col_2.metric(
                        'Суммарный охват',
                        f"{_feedback_to_float(metric_lookup.get('Суммарный охват', {}).get('После', 0.0), 0.0):,.0f}",
                        f"{_feedback_to_float(metric_lookup.get('Суммарный охват', {}).get('Изменение', 0.0), 0.0):+,.0f}"
                    )
                    metric_col_3.metric(
                        'Среднее время',
                        f"{_feedback_to_float(metric_lookup.get('Среднее время оповещения', {}).get('После', 0.0), 0.0):.2f}",
                        f"{_feedback_to_float(metric_lookup.get('Среднее время оповещения', {}).get('Изменение', 0.0), 0.0):+.2f}"
                    )
                    metric_col_4.metric(
                        'Средняя надежность',
                        f"{_feedback_to_float(metric_lookup.get('Средняя надежность', {}).get('После', 0.0), 0.0):.3f}",
                        f"{_feedback_to_float(metric_lookup.get('Средняя надежность', {}).get('Изменение', 0.0), 0.0):+.3f}"
                    )
                    st.dataframe(feedback_metrics_df, use_container_width=True, hide_index=True)

                st.markdown(f"**Итоговое решение: {st.session_state.get('feedback_decision', 'Сохранение текущей конфигурации')}**")
                st.write(st.session_state.get('feedback_decision_comment', 'Нет комментария.'))

                st.markdown('#### Изменение параметров обратной связи')
                st.dataframe(feedback_parameter_changes_df, use_container_width=True, hide_index=True)

                st.markdown('#### Сравнение решений по зонам')
                st.dataframe(feedback_comparison_df, use_container_width=True, hide_index=True)

                feedback_excel = build_feedback_excel(
                    df_before_feedback,
                    df_feedback_res,
                    feedback_comparison_df,
                    feedback_parameter_changes_df,
                    feedback_metrics_df
                )
                st.download_button(
                    label='СКАЧАТЬ РЕЗУЛЬТАТЫ ОБРАТНОЙ СВЯЗИ В EXCEL',
                    data=feedback_excel,
                    file_name='Результаты_обратной_связи.xlsx',
                    mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    use_container_width=True
                )


        # ==========================================
        # БЛОК ИИ-АНАЛИТИКИ
        # ==========================================
        st.markdown("---")
        st.subheader("ИИ-анализ результатов")
        
        if st.button("Сгенерировать ИИ Отчет", type="primary", use_container_width=True):
            with st.spinner("Аналитическая подсистема (GPT-4o) верстает официальный документ..."):
                total_cost = df_res['Стоимость'].sum()
                total_cov = df_res['Охват'].sum()
                total_pop = data_result['Население'].sum()

                # Генерация текста отчета
                ai_report = generate_ai_insights(total_cost, total_cov, total_pop, summary_table)

                # Нормализация тегов с защитой от ошибок ИИ
                ai_report_clean = ai_report.replace('[ТАБЛИЦА_ОБОРУДОВАНИЯ]', 'ТАБЛИЦА_ОБОРУДОВАНИЯ').replace('ТАБЛИЦА_ОБОРУДОВАНИЯ', '[ТАБЛИЦА_ОБОРУДОВАНИЯ]')
                ai_report_clean = ai_report_clean.replace('[ГРАФИК_РАСПРЕДЕЛЕНИЯ]', 'ГРАФИК_РАСПРЕДЕЛЕНИЯ').replace('ГРАФИК_РАСПРЕДЕЛЕНИЯ', '[ГРАФИК_РАСПРЕДЕЛЕНИЯ]')
                ai_report_clean = ai_report_clean.replace('[ГРАФИК_ОХВАТА]', 'ГРАФИК_ОХВАТА').replace('ГРАФИК_ОХВАТА', '[ГРАФИК_ОХВАТА]')

                # Генерация графиков в память
                fig1, ax1 = plt.subplots(figsize=(7, 4))
                sns.barplot(data=summary_table.sort_values('Количество', ascending=False), x='Количество', y='Система_и_Канал', palette='Blues_r', ax=ax1)
                ax1.set_title('Спецификация распределения ТСО', fontsize=11, fontweight='bold')
                ax1.set_xlabel('Количество (единиц)')
                ax1.set_ylabel('')
                plt.tight_layout()
                img1_stream = BytesIO()
                fig1.savefig(img1_stream, format='png', dpi=150)
                img1_stream.seek(0)
                plt.close(fig1)

                fig2, ax2 = plt.subplots(figsize=(7, 4))
                sns.barplot(data=summary_table.sort_values('Общий_охват', ascending=False), x='Общий_охват', y='Система_и_Канал', palette='Oranges_r', ax=ax2)
                ax2.set_title('Прогнозируемый охват населения', fontsize=11, fontweight='bold')
                ax2.set_xlabel('Охват (человек)')
                ax2.set_ylabel('')
                plt.tight_layout()
                img2_stream = BytesIO()
                fig2.savefig(img2_stream, format='png', dpi=150)
                img2_stream.seek(0)
                plt.close(fig2)

                st.success("Аналитическая записка сформирована.")
                
                # РЕНДЕРИНГ В ИНТЕРФЕЙСЕ
                with st.container(border=True):
                    parts = ai_report_clean.split('[ТАБЛИЦА_ОБОРУДОВАНИЯ]')
                    st.write(parts[0].strip())  # Заменено st.text на st.write для красивого переноса слов
                    
                    if len(parts) > 1:
                        display_df = summary_table[['Система_и_Канал', 'Количество', 'Общий_охват', 'Общая_стоимость']]
                        display_df.columns = ['Тип оборудования', 'Закупка (шт.)', 'Охват (чел.)', 'Бюджет (у.е.)']
                        st.table(display_df)
                        
                        parts2 = parts[1].split('[ГРАФИК_РАСПРЕДЕЛЕНИЯ]')
                        st.write(parts2[0].strip())
                        
                        if len(parts2) > 1:
                            st.image(img1_stream)
                            parts3 = parts2[1].split('[ГРАФИК_ОХВАТА]')
                            st.write(parts3[0].strip())
                            
                            if len(parts3) > 1:
                                st.image(img2_stream)
                                st.write(parts3[1].strip())

                # СОЗДАНИЕ ОФИЦИАЛЬНОГО DOCX ФАЙЛА
                try:
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(12)
                    
                    heading = doc.add_heading('АНАЛИТИЧЕСКАЯ ЗАПИСКА', level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    for line in ai_report_clean.split('\n'):
                        clean_line = line.strip()
                        if not clean_line: continue
                        
                        if '[ТАБЛИЦА_ОБОРУДОВАНИЯ]' in clean_line:
                            table = doc.add_table(rows=1, cols=4)
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Тип оборудования'
                            hdr_cells[1].text = 'Количество (шт.)'
                            hdr_cells[2].text = 'Охват (чел.)'
                            hdr_cells[3].text = 'Бюджет (у.е.)'
                            
                            for _, row in summary_table.iterrows():
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(row['Система_и_Канал'])
                                row_cells[1].text = str(row['Количество'])
                                row_cells[2].text = str(row['Общий_охват'])
                                row_cells[3].text = str(row['Общая_стоимость'])
                            doc.add_paragraph()
                            
                        elif '[ГРАФИК_РАСПРЕДЕЛЕНИЯ]' in clean_line:
                            doc.add_picture(img1_stream, width=Inches(6.0))
                            
                        elif '[ГРАФИК_ОХВАТА]' in clean_line:
                            doc.add_picture(img2_stream, width=Inches(6.0))
                            
                        else:
                            p = doc.add_paragraph(clean_line)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    bio = BytesIO()
                    doc.save(bio)

                    st.download_button(
                        label="Скачать официальный отчет (DOCX с таблицей и графиками)",
                        data=bio.getvalue(),
                        file_name="Analytic_Report_Official.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Ошибка формирования файла Word (DOCX): {e}")

else:
    st.info("Внимание: Отсутствуют требуемые исходные файлы данных (Excel/GeoJSON). Загрузка прервана.")
